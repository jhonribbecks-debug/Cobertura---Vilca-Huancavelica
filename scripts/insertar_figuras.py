# -*- coding: utf-8 -*-
"""Inserta:
  - 3 placeholders de modos de vibracion (Fig. 14-16) en 4.1, tras la Tabla N 11.
  - 3 esquemas de drenaje pluvial (Fig. 17-19) en el Capitulo 8.
en el documento final ya construido. Las imagenes/parrafos se crean SIEMPRE
directamente sobre `doc` (para que la relacion r:embed de cada imagen se
registre en el paquete correcto) y luego se REUBICAN (move, no copy) junto
al ancla correspondiente."""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/mnt/user-data/uploads/Tenorious/Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx"
OUT = "/tmp/Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx"
IMG_DIR = "/tmp"


def _texto(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _figura(doc, img_path, caption, width_cm=13.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = Pt(11)
    run2 = p2.add_run(caption)
    run2.font.name = 'Arial'
    run2.font.size = Pt(9)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return [p, p2]


def _mover_tras(anchor_para, paragraphs):
    """Mueve (no copia) los `paragraphs` para que queden, en orden, justo
    despues de `anchor_para`. lxml desconecta cada elemento de su padre
    actual al reinsertarlo, por lo que no hay duplicacion de relaciones."""
    prev = anchor_para._p
    for p in paragraphs:
        prev.addnext(p._p)
        prev = p._p


def main():
    doc = Document(SRC)

    # ------------------------------------------------------------------
    # 1) Placeholders de modos de vibracion (Fig. 14-16), tras Tabla N 11
    # ------------------------------------------------------------------
    anchor_modal = None
    for p in doc.paragraphs:
        if p.text.strip() == "Fuente: SAP2000 (análisis modal).":
            anchor_modal = p
            break
    assert anchor_modal is not None, "No se encontro el anchor de la Tabla N 11"

    nuevos = []
    nuevos.append(_texto(doc,
        "Las Figuras N° 14 a 16 ilustran las deformadas de los tres modos "
        "de vibración con mayor masa participativa en cada dirección "
        "principal (X, Y, Z). El espacio de cada figura queda reservado "
        "para la captura de pantalla de SAP2000 (vista isométrica, con "
        "indicación de la escala de deformación empleada), que se "
        "incorporará en una revisión posterior del documento."))
    modos = [
        ("esq_modo1.png",
         "Figura N° 14.- Deformada del Modo N° 1 (T₁ = 0.6252 s; traslación "
         "predominante en X, UX = 93.60 %). Espacio reservado para la "
         "captura de SAP2000; pendiente de incorporar por el proyectista."),
        ("esq_modo2.png",
         "Figura N° 15.- Deformada del Modo N° 2 (T₂ = 0.4895 s; traslación "
         "predominante en Y, UY = 92.44 %). Espacio reservado para la "
         "captura de SAP2000; pendiente de incorporar por el proyectista."),
        ("esq_modo4.png",
         "Figura N° 16.- Deformada del Modo N° 4 (T₄ = 0.3395 s; traslación "
         "predominante en Z / vertical, UZ = 46.59 %). Espacio reservado "
         "para la captura de SAP2000; pendiente de incorporar por el "
         "proyectista."),
    ]
    for fname, cap in modos:
        nuevos.extend(_figura(doc, os.path.join(IMG_DIR, fname), cap))
    _mover_tras(anchor_modal, nuevos)
    print("  + Figuras 14-16 (modos de vibración, placeholders) insertadas en 4.1")

    # ------------------------------------------------------------------
    # 2) Esquemas de drenaje pluvial (Fig. 17-19) en Capitulo 8
    # ------------------------------------------------------------------
    anchor_planta = anchor_canaleta = anchor_bajante = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Tabla N° 26.- Reparto de caudal"):
            anchor_planta = p
        elif t.startswith("Q capacidad (35.4 L/s)"):
            anchor_canaleta = p
        elif t.startswith("Q capacidad (10.59 L/s)"):
            anchor_bajante = p
    assert anchor_planta and anchor_canaleta and anchor_bajante, "Faltan anchors de drenaje"

    fig17 = _figura(doc, os.path.join(IMG_DIR, "esq_planta_drenaje.png"),
        "Figura N° 17.- Esquema en planta del área de captación, canalones y "
        "bajantes pluviales (esquemático, no a escala de replanteo).")
    _mover_tras(anchor_planta, fig17)
    print("  + Figura 17 (planta de drenaje) insertada en 8.5")

    fig18 = _figura(doc, os.path.join(IMG_DIR, "esq_canaleta.png"),
        "Figura N° 18.- Sección transversal esquemática de la canaleta "
        "metálica tipo caja (0.25 × 0.14 m), con el tirante de diseño "
        "(y = 0.112 m) y la pendiente longitudinal (S = 1.0 %).")
    _mover_tras(anchor_canaleta, fig18)
    print("  + Figura 18 (sección canaleta) insertada en 8.6")

    fig19 = _figura(doc, os.path.join(IMG_DIR, "esq_bajante.png"),
        "Figura N° 19.- Detalle esquemático de la conexión canalón-bajante "
        "pluvial (PVC Ø 2\", n = 0.010).")
    _mover_tras(anchor_bajante, fig19)
    print("  + Figura 19 (detalle bajante) insertada en 8.7")

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == "__main__":
    main()
