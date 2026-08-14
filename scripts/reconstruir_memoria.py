"""RECONSTRUCCIÓN COMPLETA — Memoria de Cálculo Estructural Cobertura Huancalpi.

Estructura final (orden PRONIED + Drenaje como Cap. 8):

    PORTADA (con logotipo de la empresa)
    ÍNDICE GENERAL (campo TOC, se actualiza con Word)
    Índice de Tablas y Figuras (nota)
    1.  GENERALES
    2.  ANÁLISIS POR CARGAS DE GRAVEDAD
    3.  ANÁLISIS SÍSMICO Y DE VIENTO
    4.  RESULTADOS DEL ANÁLISIS
    5.  DISEÑO DE ELEMENTOS ESTRUCTURALES
        5.4  CONEXIONES, PLACAS BASE Y PERNOS DE ANCLAJE  (completo:
             datos reales del informe IDEA StatiCa + páginas del PDF como figuras)
    6.  CIMENTACIÓN
    7.  CONCLUSIONES Y COMENTARIOS
    8.  DRENAJE PLUVIAL   (integrado como capítulo, sin "texto íntegro")
    ANEXO A. PLANOS
    FIRMAS

Se eliminan: índice estático duplicado, "CONTROL DE VERSIONES" mal ubicado,
la figura isométrica "Vista 3D (Revit)" del anexo.

FORMATO PRONIED:
    Arial, texto justificado, H1/H2/H3 Arial #1F497D/#2E75B6,
    tablas "Tabla N° X" + fuente, encabezado #1F497D, figuras "Figura N° X"
    con leyenda cursiva, numeración de páginas en el pie.

Ejecutar despues: actualizar_indice_y_pdf.py (actualiza TOC y exporta PDF).
"""

from __future__ import annotations

import copy
import os
import re
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(ROOT, "Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.original.docx")
OUT = os.path.join(ROOT, "Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx")
CONN_PDF = os.path.join(ROOT, "Informes de conexión metálica columna PedestaL.pdf")
FIG_DIR = os.path.join(ROOT, "_conn_figuras")
LOGO = os.path.join(ROOT, "Gemini_Generated_Image_1iyknz1iyknz1iyk-removebg-preview.png")

AZUL_N1 = RGBColor(0x1F, 0x49, 0x7D)
AZUL_N2 = RGBColor(0x2E, 0x75, 0xB6)
FONDO_TABLA = "1F497D"
FONDO_CELDA = "D9E2F3"

# Índice estático duplicado del original (se elimina; el TOC field queda)
_STATIC_INDICE = [
    "1.  GENERALES", "1.1  Alcance del documento", "1.2  Objetivos",
    "1.3  Descripción del sistema estructural",
    "1.4  Bases legales y normativa aplicable",
    "1.5  Información general del proyecto", "1.6  Materiales y perfiles",
    "1.7  Condiciones de cimentación",
    "2.  ANÁLISIS POR CARGAS DE GRAVEDAD", "2.1  Modelo estructural",
    "2.2  Secciones empleadas", "2.3  Metrado de cargas",
    "2.4  Combinaciones de carga",
    "3.  ANÁLISIS SÍSMICO Y DE VIENTO",
    "3.1  Parámetros sísmicos (NTP E.030)",
    "3.2  Periodo fundamental y masas participativas",
    "3.3  Cortante basal estático (verificación)",
    "3.4  Cargas de viento (NTP E.020)", "3.5  Control de desplazamientos",
    "4.  RESULTADOS DEL ANÁLISIS", "4.1  Periodos y modos de vibración",
    "4.2  Desplazamientos máximos", "4.3  Reacciones en apoyos",
    "4.4  Fuerzas internas envolventes por sección",
    "5.  DISEÑO DE ELEMENTOS ESTRUCTURALES", "5.1  Criterios de diseño",
    "5.2  Memoria de diseño de elementos (fórmulas y verificación)",
    "5.3  Resumen de verificación (D/C Ratio)",
    "5.4  Conexiones, placas base y pernos de anclaje",
    "5.5  Cimentación",
    "6.  CONCLUSIONES Y COMENTARIOS", "ANEXO A  PLANOS",
]


# ------------------------------------------------------------- ESTILOS

def _configurar_pagina(doc: Document):
    """A4 vertical con los márgenes del documento original PRONIED."""
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(3.2)
    sec.bottom_margin = Cm(2.2)


def _configurar_estilos(doc: Document):
    for style in doc.styles:
        try:
            if style.type == 1:
                style.font.name = 'Arial'
        except Exception:
            pass

    s1 = doc.styles['Heading 1']
    s1.font.name = 'Arial'
    s1.font.size = Pt(14)
    s1.font.bold = True
    s1.font.color.rgb = AZUL_N1
    s1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s1.paragraph_format.space_before = Pt(18)
    s1.paragraph_format.space_after = Pt(6)
    s1.paragraph_format.line_spacing = Pt(16)

    s2 = doc.styles['Heading 2']
    s2.font.name = 'Arial'
    s2.font.size = Pt(12)
    s2.font.bold = True
    s2.font.color.rgb = AZUL_N2
    s2.paragraph_format.space_before = Pt(12)
    s2.paragraph_format.space_after = Pt(4)

    s3 = doc.styles['Heading 3']
    s3.font.name = 'Arial'
    s3.font.size = Pt(11)
    s3.font.bold = True
    s3.font.color.rgb = AZUL_N2

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.5)

    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Arial'
        lb.font.size = Pt(10)
        lb.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        lb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lb.paragraph_format.line_spacing = Pt(12)
        lb.paragraph_format.left_indent = Cm(1.0)
        lb.paragraph_format.first_line_indent = Cm(-0.5)


def _pie_de_pagina(doc: Document):
    """Numeración 'Página X de Y' centrada en el pie (Arial Narrow 8.5pt)."""
    sec = doc.sections[0]
    foot = sec.footer
    foot.is_linked_to_previous = False
    p = foot.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    def run(txt=None, rpr=True):
        r = p.add_run()
        if rpr:
            r.font.name = 'Arial Narrow'
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if txt is not None:
            r.text = txt
        return r

    def field(instr):
        r1 = p.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        r1._r.append(fld)
        r2 = p.add_run()
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = instr
        r2._r.append(it)
        r3 = p.add_run()
        sep = OxmlElement('w:fldChar')
        sep.set(qn('w:fldCharType'), 'separate')
        r3._r.append(sep)
        r4 = p.add_run('1')
        r5 = p.add_run()
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        r5._r.append(end)
        for r in (r1, r2, r3, r4, r5):
            r.font.name = 'Arial Narrow'
            r.font.size = Pt(8.5)

    run('Página ')
    field('PAGE')
    run(' de ')
    field('NUMPAGES')
    run('   |   Elaborado por: Ing. Jhon Brian Ribbeck Soto - C.I.P. 289452')


# ------------------------------------------------------------- HELPERS

def _h(doc, text, level):
    return doc.add_paragraph(text, style=f"Heading {level}")


def _para(doc, text, size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def _bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = size
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def _table(doc, headers, rows, caption=None, size=Pt(8.5)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), FONDO_TABLA)
        shading.set(qn('w:val'), 'clear')
        hdr[j]._tc.get_or_add_tcPr().append(shading)

    for i, r in enumerate(rows):
        row = t.add_row()
        for j, v in enumerate(r):
            row.cells[j].text = "" if v is None else str(v)
            for par in row.cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), FONDO_CELDA)
                shading.set(qn('w:val'), 'clear')
                row.cells[j]._tc.get_or_add_tcPr().append(shading)

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return t


def _figura(doc, img_path, caption, width_cm=13.5):
    if not os.path.exists(img_path):
        _para(doc, f"[FALTA IMAGEN: {os.path.basename(img_path)}]", size=Pt(9))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(caption)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


# ----------------------------------------------------- CONTENIDO 5.4 CONEXIONES

def _crear_54_conexiones(doc):
    _h(doc, "5.4 Conexiones, placas base y pernos de anclaje", 2)
    _para(doc, "La placa base de las columnas se verificó con el modelo de "
               "elementos finitos «Conección plancha base» (IDEA StatiCa, "
               "método CBFEM); el informe completo se adjunta como archivo "
               f"'{os.path.basename(CONN_PDF)}'. Las conexiones se diseñan "
               "soldadas (AWS D1.1) o empernadas (ASTM A325). La verificación "
               "se realiza bajo las combinaciones de carga del Capítulo 2.")

    _h(doc, "5.4.1 Datos de diseño y materiales", 3)
    _table(doc, ["Concepto", "Valor", "Unidad"], [
        ["Proyecto", "HUANCALPI-RIBBECK", "—"],
        ["Norma estructural", "AISC 360-16 (LRFD)", "—"],
        ["Norma de anclajes", "ACI 318-14 §17.4", "—"],
        ["Sección transversal columna", "4 - HSS(Imp) 8×8×5/8", "—"],
        ["Material columna", "A36 (Fy = 36 ksi; εlim = 5 %)", "—"],
        ["Perno de anclaje", "¾\" A325", "fu = 119.7 ksi"],
        ["Diámetro anclaje", "1.905", "cm"],
        ["Área bruta del anclaje", "2.850", "cm²"],
        ["Longitud de anclaje", "100.0", "cm"],
        ["Transferencia de corte", "Fricción / junta de mortero (2.40 cm)", "—"],
        ["Hormigón de apoyo", "4000 / 3000", "psi"],
    ], caption="Tabla N° 17.- Datos de diseño y materiales de la conexión.")

    _h(doc, "5.4.2 Cargas de equilibrio (LE1)", 3)
    _para(doc, "Fuerzas de equilibrio (LE1) en el Joint COL de la placa base, "
               "obtenidas del análisis global del modelo SAP2000:")
    _table(doc, ["Fuerza", "Valor", "Unidad"], [
        ["N (compresión)", "-66.539", "kN"],
        ["Vy", "4.901", "kN"],
        ["Vz", "-29.184", "kN"],
        ["Mx", "0.00", "kN·m"],
        ["My", "-14.03", "kN·m"],
        ["Mz", "-85.66", "kN·m"],
    ], caption="Tabla N° 18.- Cargas de equilibrio en la columna.")
    _figura(doc, os.path.join(FIG_DIR, "fig_col.png"),
            "Figura N° 6.- Elemento COL: sección HSS(Imp) 8×8×5/8, orientación "
            "e inserción en el modelo CBFEM (página 2 del informe).")

    _h(doc, "5.4.3 Resumen de verificaciones", 3)
    _table(doc, ["Concepto", "Ratio", "Límite", "Estado"], [
        ["Análisis general", "100.0", "%", "OK"],
        ["Placas (BP1)", "0.7", "< 5.0 %", "OK"],
        ["Anclajes (tracción)", "99.3", "< 100 %", "OK"],
        ["Soldaduras", "88.4", "< 100 %", "OK"],
        ["Bloque de hormigón", "12.9", "< 100 %", "OK"],
        ["Corte bloque-hormigón", "36.7", "< 100 %", "OK"],
        ["Pandeo", "—", "No calculado", "—"],
    ], caption="Tabla N° 19.- Resumen de estados de verificación (LE1).")
    _figura(doc, os.path.join(FIG_DIR, "fig_verif.png"),
            "Figura N° 7.- Verificación general y deformación de la conexión "
            "para LE1 (página 6 del informe IDEA StatiCa).")

    _h(doc, "5.4.4 Verificación de placas por elemento", 3)
    _table(doc, ["Elemento", "Fy (ksi)", "Espesor", "σEd (ksi)",
                 "εPl (%)", "σcEd (ksi)", "Estado"], [
        ["COL", "36.0", "5/8\"", "32.4", "0.0", "0.0", "OK"],
        ["BP1 (placa base)", "36.0", "1\"", "28.5", "0.0", "0.0", "OK"],
        ["RIB2a", "36.0", "1/4\"", "32.4", "0.1", "0.0", "OK"],
        ["RIB2b", "36.0", "1/4\"", "32.5", "0.3", "0.0", "OK"],
        ["RIB4a", "36.0", "1/4\"", "32.5", "0.2", "0.0", "OK"],
        ["RIB4b", "36.0", "1/4\"", "32.6", "0.7", "0.0", "OK"],
    ], caption="Tabla N° 20.- Tensión/compresión por elemento (LE1).")
    _figura(doc, os.path.join(FIG_DIR, "fig_placas.png"),
            "Figura N° 8.- Bloque CB1 y verificación de placas BP1/RIB "
            "(página 5 del informe).")

    _h(doc, "5.4.5 Verificación de anclajes", 3)
    _table(doc, ["Concepto", "ϕRn (kN)", "≥ Demanda (kN)", "Norma", "Estado"], [
        ["Tracción anclaje (A2)", "124.440", "84.262",
         "ACI 318-14 §17.4.1 (ϕ=0.70)", "OK"],
        ["Arrancamiento tracción grupo (A2,A3,A7)", "239.042", "237.345",
         "ACI 318-14 §17.4.2", "OK"],
        ["Arrancamiento lateral grupo", "984.889", "84.262",
         "ACI 318-14 §17.4.4 (ψ=0.51)", "OK"],
        ["Arrancamiento barra (A2-A7)", "1125.290", "0.0",
         "ACI 318-14 §17.4.3", "OK"],
        ["Pryout del hormigón (grupo)", "580.139", "0.0",
         "ACI 318-14 §17.5.3", "OK"],
        ["Aplastamiento bloque-hormigón", "5.1 ksi", "0.4 ksi",
         "AISC J8", "OK"],
        ["Interacción tracción-cortante (ACI 17.6)", "0.99", "≤ 1.0",
         "ACI 318-14 §R17.6", "OK"],
    ], caption="Tabla N° 21.- Verificación de anclajes a tracción/lateral.")
    _para(doc, "Fórmula de anclaje a tracción (ACI 318-14 §17.4.1): "
               "ϕNsa = ϕ·Asa·fu = 0.70 × 2.155 cm² × 119.7 ksi = 124.440 kN "
               "≥ 84.262 kN. Cono de arrancamiento (§17.4.2): Acon = "
               "12 636.5 cm²; Aco (individual) = 9 409 cm²; Ψed,N = 0.98. "
               "Arrancamiento lateral (§17.4.4): Área portante de la cabeza "
               "= 97 150 cm²; ψ = 0.51 (ca1 = 46 cm, ca2 = 48.5 cm, s = 15 cm). "
               "Pryout (§17.5.3): ϕ = 0.65; kcp = 2.0; ϕVcp = 580.139 kN.",
               size=Pt(8.5))
    _figura(doc, os.path.join(FIG_DIR, "fig_anclajes.png"),
            "Figura N° 9.- Verificación del anclaje A2 a tracción "
            "(página 7 del informe).")

    _h(doc, "5.4.6 Verificación de soldaduras", 3)
    _table(doc, ["Conexión", "ϕRn (kN)", "≥ Demanda (kN)", "Áng. (°)",
                 "Estado"], [
        ["BP1 / RIB2a", "11.392", "8.649", "83.8", "OK"],
        ["COL-w2 / RIB2a", "10.008", "7.630", "46.5", "OK"],
        ["BP1 / RIB2b", "11.358", "8.927", "81.1", "OK"],
        ["COL-w2 / RIB2b", "10.016", "7.632", "46.6", "OK"],
        ["BP1 / RIB4a", "11.369", "8.913", "81.9", "OK"],
        ["COL-w4 / RIB4a", "10.050", "7.689", "47.2", "OK"],
        ["COL-w4 / RIB4b", "10.089", "7.835", "47.8", "OK"],
        ["BP1 / COL", "6.687", "5.910", "79.1", "OK"],
    ], caption="Tabla N° 22.- Resistencia de soldaduras (ϕ = 0.75, E70xx "
               "70 ksi) según AISC 360-16 §J2.4 / AWS D1.1.")
    _figura(doc, os.path.join(FIG_DIR, "fig_soldaduras.png"),
            "Figura N° 10.- Verificación de soldaduras BP1/RIB2a "
            "(página 17 del informe).")

    _h(doc, "5.4.7 Verificación del bloque de cimentación", 3)
    _table(doc, ["Concepto", "ϕRn", "≥ Demanda", "Norma", "Estado"], [
        ["Compresión bloque (BP1)", "5.1 ksi", "0.4 ksi", "AISC J8", "OK"],
        ["Corte bloque-hormigón", "—", "—", "ACI 318-14 (ϕVcp)", "OK"],
    ], caption="Tabla N° 23.- Verificación del bloque de cimentación CB1.")
    _para(doc, "Bloque CB1: 127.000 × 127.000 cm; profundidad 150.0 cm; "
               "longitud de anclaje 100.0 cm; área de contacto 1049.627 cm²; "
               "superficie de apoyo 12 734.202 cm²; ϕ = 0.65; transferencia "
               "de corte por fricción / junta de mortero (2.40 cm).",
               size=Pt(8.5))
    _figura(doc, os.path.join(FIG_DIR, "fig_material.png"),
            "Figura N° 11.- Listado de material, operaciones, soldaduras y "
            "anclajes; dibujo de la placa base BP1 (página 24 del informe).")

    _h(doc, "5.4.8 Planos de fabricación y listado de material", 3)
    _table(doc, ["Elemento", "Material", "Cant.", "Soldadura", "Observación"], [
        ["BP1 — placa base 50×50 cm, e = 1\"", "A36", "1",
         "simple a = 1/8\" (72.3 cm)", "P1\"×50-50"],
        ["RIB2 — cartela 10×20 cm, e = 1/4\"", "A36", "2",
         "doble a = 1/8\" (60.0 cm)", "P1/4\"×10-20"],
        ["RIB4 — cartela 10×20 cm, e = 1/4\"", "A36", "2",
         "doble a = 1/8\" (60.0 cm)", "P1/4\"×10-20"],
        ["Anclajes ¾\" A325", "A325", "6", "—",
         "L = 105 cm; taladro 100 cm"],
    ], caption="Tabla N° 24.- Listado de material de la conexión.")
    _figura(doc, os.path.join(FIG_DIR, "fig_bp1_rib2.png"),
            "Figura N° 12.- Planos de fabricación: placa base BP1 (P1\"×50-50) "
            "y cartelas RIB2 (página 25 del informe).")
    _figura(doc, os.path.join(FIG_DIR, "fig_rib4.png"),
            "Figura N° 13.- Plano de fabricación: cartelas RIB4 y explicación "
            "de símbolos de soldadura (página 26 del informe).")


# ------------------------------------------------------------- CONTENIDO 8 DRENAJE

def _crear_8_drenaje(doc):
    _h(doc, "8. DRENAJE PLUVIAL", 1)
    _para(doc, "Este capítulo integra el informe de drenaje pluvial de la "
               "cobertura metálica en arco, elaborado conforme a la Norma "
               "Técnica CE.040 «Drenaje Pluvial» del Reglamento Nacional de "
               "Edificaciones (RNE). Unidades: m², mm/h, L/s, m.")

    _h(doc, "8.1 Objeto y alcance", 2)
    _para(doc, "Determinar el caudal de diseño de la escorrentía pluvial "
               "generada sobre la cobertura y dimensionar/verificar las "
               "canaletas y bajantes pluviales que evacúan el agua hacia el "
               "sistema de drenaje del proyecto. El alcance se limita al "
               "cálculo hidrológico-hidráulico de la instalación de drenaje "
               "pluvial de la cobertura (área captante).")

    _h(doc, "8.2 Normativa aplicable", 2)
    _bullet(doc, "CE.040 «Drenaje Pluvial» (RNE, RM 126-2021): Art. 11 "
                 "(caudal de diseño), Art. 12.4 (canaletas), Art. 12.5 "
                 "(bajantes).")
    _bullet(doc, "CE.040 — Anexo I (Hidrología): método racional, coeficiente "
                 "de escorrentía, parámetros pluviométricos.")
    _bullet(doc, "CE.040 — Anexo II (Hidráulica): ecuación de Manning, "
                 "coeficientes de rugosidad.")

    _h(doc, "8.3 Descripción del proyecto", 2)
    _para(doc, "El proyecto se ubica en el distrito de Vilca, provincia de "
               "Huancavelica, departamento de Huancavelica, a una altitud de "
               "3 275 msnm. La cobertura es metálica en arco (bóveda), con "
               "plancha de acero recubierta Aluzinc, perfil TR5, recubrimiento "
               "AZ150, y área en planta (proyección horizontal) de 80.295 m².")

    _h(doc, "8.4 Variables y parámetros de diseño", 2)
    _table(doc, ["Variable", "Símbolo", "Valor adoptado", "Fuente normativa"], [
        ["Área de techo (proyección horizontal)", "A", "80.295 m²",
         "CE.040, Anexo I, num. 1.2.4"],
        ["Coeficiente de escorrentía", "C", "0.95",
         "CE.040, Anexo I, Tabla 1.b — cobertura metálica"],
        ["Periodo de retorno", "T", "25 años",
         "CE.040, Art. 11.1 (criterio del proyectista)"],
        ["Tiempo de concentración / duración", "tc", "10 min = 0.1667 h",
         "CE.040, Anexo I, num. 1.2.3"],
        ["Intensidad de precipitación de diseño", "I", "30.33 mm/h",
         "CE.040, Anexo I, num. 1.3 (IILA-SENAMHI-UNI)"],
    ], caption="Tabla N° 25.- Cuadro resumen de variables de diseño.")

    _h(doc, "8.5 Cálculo del caudal de diseño", 2)
    _para(doc, "El caudal se calcula con el método racional "
               "(CE.040, Art. 11.2/11.3 y Anexo I, num. 1.2.1): "
               "Q = C·I·A/3600. Reemplazando: Q = 0.95 × 30.33 × 80.295/3600 "
               "= 0.643 L/s (caudal de diseño total de la cobertura).")
    _table(doc, ["Elemento", "N°", "Área tributaria", "Caudal"], [
        ["Cobertura (total)", "1", "80.295 m²", "0.643 L/s"],
        ["Canalón", "2", "40.148 m² c/u", "0.321 L/s c/u"],
        ["Bajante pluvial", "4 (2 por canalón)", "20.074 m² c/u",
         "0.161 L/s c/u"],
    ], caption="Tabla N° 26.- Reparto de caudal entre canalones y bajantes.")

    _h(doc, "8.6 Verificación hidráulica de canaletas", 2)
    _para(doc, "La CE.040 (Art. 12.4) exige canaletas calculadas para un "
               "eficaz discurrimiento del agua. Se adopta una canaleta "
               "metálica tipo caja (rectangular) de 0.25 × 0.14 m según el "
               "detalle constructivo del alero, con pendiente S = 1 %. "
               "Verificación con la ecuación de Manning (n = 0.013):")
    _table(doc, ["Parámetro", "Valor", "Fuente"], [
        ["Sección adoptada", "Rectangular (caja), 0.25 m × 0.14 m",
         "Detalle constructivo del proyecto"],
        ["Tirante de diseño", "y = 80% de la altura = 0.112 m",
         "Criterio conservador de verificación"],
        ["Pendiente", "S = 1 % (mínimo normativo)", "CE.040, Art. 12.4"],
        ["Capacidad (y = 0.112 m)", "Q ≈ 35.4 L/s", "Ecuación de Manning"],
        ["Capacidad (sección llena)", "Q ≈ 47.6 L/s", "Referencial, sin resguardo"],
    ], caption="Tabla N° 27.- Verificación hidráulica de la canaleta.")
    _para(doc, "Q capacidad (35.4 L/s) ≥ Q diseño por canalón (0.321 L/s) "
               "→ VERIFICACIÓN CONFORME.")

    _h(doc, "8.7 Verificación y disposición de bajantes pluviales", 2)
    _para(doc, "La CE.040 (Art. 12.5) exige un diámetro mínimo de bajante de "
               "0.05 m (2\"). Se verifica su capacidad con la ecuación de "
               "Manning para sección circular llena (PVC, n = 0.010, tramo "
               "vertical asimilado a S = 1):")
    _table(doc, ["Parámetro", "Valor", "Fuente"], [
        ["Diámetro adoptado", "D = 0.05 m (2\") — mínimo normativo",
         "CE.040, Art. 12.5"],
        ["Coeficiente de rugosidad (PVC)", "n = 0.010",
         "CE.040, Anexo II, Gráfico 1"],
        ["Pendiente asimilada", "S = 1 (tramo vertical)",
         "Criterio conservador de verificación"],
        ["Capacidad", "Q ≈ 10.59 L/s", "Ecuación de Manning"],
    ], caption="Tabla N° 28.- Verificación de la bajante pluvial.")
    _para(doc, "Q capacidad (10.59 L/s) ≥ Q diseño por bajante (0.161 L/s) "
               "→ VERIFICACIÓN CONFORME. Hidráulicamente, un solo bajante por "
               "canalón sería suficiente; la disposición adoptada de 2 "
               "bajantes por canalón responde a criterio constructivo y "
               "redundancia del proyectista. Espaciamiento entre bajantes: "
               "30.30 m / 2 = 15.15 m.")

    _h(doc, "8.8 Conclusiones", 2)
    _bullet(doc, "El caudal de diseño de la cobertura, por el método racional "
                 "de la CE.040 con C = 0.95, I = 30.33 mm/h y A = 80.295 m², "
                 "es Q = 0.643 L/s.")
    _bullet(doc, "Repartido entre los dos canalones y los cuatro bajantes "
                 "(2 por canalón), el caudal individual por bajante es "
                 "0.161 L/s.")
    _bullet(doc, "La canaleta tipo caja de 0.25 × 0.14 m (pendiente 1 %) y "
                 "los bajantes Ø50 mm cumplen holgadamente las exigencias de "
                 "la CE.040.")
    _bullet(doc, "Los valores de periodo de retorno (T = 25 años) y de "
                 "subzona pluviométrica (I23₁₀, estación 757, Huancavelica) "
                 "son aproximaciones a partir de la información pública "
                 "disponible (SENAMHI); deben confirmarse con los registros "
                 "oficiales antes de la fabricación si se requiere mayor "
                 "precisión.")

    _h(doc, "8.9 Referencias normativas", 2)
    _bullet(doc, "Ministerio de Vivienda, Construcción y Saneamiento. Norma "
                 "Técnica CE.040 «Drenaje Pluvial» del RNE.")
    _bullet(doc, "IILA-SENAMHI-UNI (1983). «Estudio de la Hidrología del "
                 "Perú».")
    _bullet(doc, "SENAMHI. Curvas Intensidad-Duración-Frecuencia (registros "
                 "de la estación de Huancavelica).")


# ------------------------------------------------------------- COPIA DEL ORIGINAL

def _append_de_origen(doc_dst, doc_src, desde, hasta=None, skip_texts=None):
    """Copia al cuerpo de doc_dst los elementos de doc_src entre el párrafo
    cuyo texto comienza con `desde` (inclusive) y el que comienza con `hasta`
    (exclusivo). Se omiten los párrafos cuyo texto coincida con skip_texts.
    Los elementos se insertan ANTES del sectPr final del destino (igual que
    add_paragraph), para mantener el orden del cuerpo."""
    body_src = doc_src.element.body
    body_dst = doc_dst.element.body
    sect_pr = body_dst.find(qn('w:sectPr'))
    started = False
    n = 0
    for el in body_src.iterchildren():
        tag = el.tag
        txt = None
        if tag == qn('w:p'):
            txt = ''.join(el.itertext()).strip()
            if not started:
                if desde is not None and txt.startswith(desde):
                    started = True
                else:
                    continue
            if hasta is not None and txt.startswith(hasta):
                break
            if skip_texts and txt in skip_texts:
                continue
        elif tag == qn('w:tbl'):
            if not started:
                continue
        elif tag == qn('w:sectPr'):
            continue
        if sect_pr is not None:
            sect_pr.addprevious(copy.deepcopy(el))
        else:
            body_dst.append(copy.deepcopy(el))
        n += 1
    return n


def _limpiar_indice_estatico(doc):
    """Elimina el índice estático duplicado (después del campo TOC y antes de
    'Índice de Tablas y Figuras')."""
    paras = doc.paragraphs
    start = None
    end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if start is None and t.startswith("Índice — en Word"):
            start = i
        elif start is not None and t.startswith("Índice de Tablas y Figuras"):
            end = i
            break
    if start is None or end is None:
        return 0
    n = 0
    for i in range(start + 1, end):
        paras[i]._p.getparent().remove(paras[i]._p)
        n += 1
    return n


def _retext_py(p_elemento, nuevo_texto):
    """Reemplaza el texto de un <w:p> conservando el primer run (formato)."""
    runs = p_elemento.findall(qn('w:r'))
    if runs:
        t_el = runs[0].find(qn('w:t'))
        if t_el is None:
            t_el = OxmlElement('w:t')
            runs[0].append(t_el)
        t_el.text = nuevo_texto
        t_el.set(qn('xml:space'), 'preserve')
        for r in runs[1:]:
            r.getparent().remove(r)
    return p_elemento


def _quitar_figura(doc, etiqueta):
    """Elimina el párrafo con `etiqueta` y el párrafo de imagen inmediato anterior."""
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if etiqueta in p.text:
            img = paras[i - 1] if i > 0 else None
            if img is not None and img._p.findall('.//' + qn('w:drawing')):
                img._p.getparent().remove(img._p)
            p._p.getparent().remove(p._p)
            return True
    return False


def _renombrar_h1(doc, texto_viejo, texto_nuevo):
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == texto_viejo:
            _retext_py(p._p, texto_nuevo)
            return True
    return False


# ------------------------------------------------------------- BUILD

def build():
    if os.path.exists(OUT):
        os.remove(OUT)

    doc = Document()
    doc.add_paragraph("")

    src = Document(SRC)
    _configurar_pagina(doc)
    _configurar_estilos(doc)
    _pie_de_pagina(doc)

    # 0. Logotipo de la empresa al inicio de la portada
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(6)
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.line_spacing = 1.0
    logo_p.paragraph_format.line_spacing_rule = 1
    if os.path.exists(LOGO):
        logo_run = logo_p.add_run()
        logo_run.add_picture(LOGO, width=Cm(2.2))
    body = doc._element.body
    first_p = body.find(qn('w:p'))
    if first_p is not None:
        logo_p._p.addprevious(first_p)
        body.remove(first_p)

    # 1. Portada + ÍNDICE GENERAL (TOC) + nota (hasta cap. 1)
    _append_de_origen(doc, src, desde="MINEDU — PRONIED", hasta="1. GENERALES")
    n = _limpiar_indice_estatico(doc)
    print(f"  + Portada + TOC copiados (índice estático suprimido: {n} renglones)")

    # 2. Capítulos 1-4 y 5.1 (hasta 5.2, que se genera con datos del .s2k v4)
    _append_de_origen(doc, src, desde="1. GENERALES",
                      hasta="5.2 Memoria de diseño de elementos")
    print("  + Capítulos 1-4 y 5.1 copiados")

    # 2a. Actualizar tablas 4/5/14 con los datos reales del modelo v4
    import cap52_53
    cap52_53.corregir_tablas_v4(doc)

    # 2b. Capítulo 5.2 Memoria de diseño + 5.3 Resumen (9 secciones reales)
    cap52_53.crear_52_53(doc)

    # 3. Capítulo 5.4 CONEXIONES completo
    _crear_54_conexiones(doc)
    print("  + 5.4 Conexiones creado (datos reales + figuras del PDF)")

    # 4. Capítulo 5.5 Cimentación -> se promueve a Capítulo 6 CIMENTACIÓN
    _append_de_origen(doc, src, desde="5.5 Cimentación",
                      hasta="6. CONCLUSIONES")
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip() == "5.5 Cimentación":
            p.style = doc.styles["Heading 1"]
            _retext_py(p._p, "6. CIMENTACIÓN")
            break
    print("  + 5.5 Cimentación promovido a Capítulo 6 CIMENTACIÓN")

    # 5. Capítulo 7 CONCLUSIONES (era "6. CONCLUSIONES Y COMENTARIOS")
    _append_de_origen(doc, src, desde="6. CONCLUSIONES Y COMENTARIOS",
                      hasta="ANEXO A")
    _renombrar_h1(doc, "6. CONCLUSIONES Y COMENTARIOS",
                  "7. CONCLUSIONES Y COMENTARIOS")
    print("  + Conclusiones renumerado a Capítulo 7")

    # 6. Capítulo 8 DRENAJE PLUVIAL
    _crear_8_drenaje(doc)
    print("  + Capítulo 8 Drenaje integrado")

    # 7. ANEXO A. PLANOS (sin la figura isométrica de Revit)
    _append_de_origen(doc, src, desde="ANEXO A. PLANOS", hasta="CONTROL DE VERSIONES")
    _quitar_figura(doc, "Figura N° 6.- Vista 3D")
    print("  + ANEXO A copiado (figura isométrica eliminada)")

    # 8. FIRMAS
    _append_de_origen(doc, src, desde="FIRMAS", hasta=None)
    print("  + FIRMAS copiadas")

    doc.save(OUT)
    print(f"OK: documento final -> {OUT}")
    return OUT


if __name__ == "__main__":
    build()
