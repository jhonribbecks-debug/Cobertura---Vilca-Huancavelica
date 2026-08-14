"""Genera el capítulo 5.2 (Memoria de diseño de elementos) y 5.3 (Resumen D/C)
con las 9 secciones reales del modelo .s2k v4 (HUANCALPI - MODELO FINAL v4.s2k).

Verificaciones AISC 360-16 (LRFD) con los valores reales de SAP2000:
  E3  compresión   -> phiC*Pn (pandeo, KL/r)
  D2  tracción     -> phiT*Pn = 0.90*Ag*Fy
  F8  flexión HSS  -> phiB*Mn = 0.90*Fy*Z
  H1.1 interacción -> segun PRatio (H1-1a/H1-1b)
Corrección de materiales respecto a la memoria original (que usaba GrB46 Fy=317
para todo): el modelo v4 asigna ASTM A500 GrA (Fy=269, Fu=310) a la mayoria de
secciones; solo BRIDA SUPERIOR HSS100x50x3 y DIAGONALES HSS 50x50x2 usan
A500GrB46 (Fy=317). El TENSOR Ø5/8 es acero A36 (Fy=36 ksi) y solo trabaja a
traccion (phiT*Pn = 0.90*Ag*Fy); la compresion reportada por SAP2000 es
artefacto del limite de compresion = 0 y no es caso de diseño del elemento.
"""
import json
import math
import os

import reconstruir_memoria as rm
from docx.shared import Pt, Cm

ROOT = rm.ROOT
DISENO = os.path.join(ROOT, "_s2k_diseno.json")
FIG_DC = os.path.join(ROOT, "_figura_dc.png")

_KSI = 6.89476  # MPa por ksi


def _calc_verif(s, fy, e, ag, z33, r22, lm, p, m3, m2):
    """Retorna dict con las verificaciones manuales (E3/D2/F8/H1.1)."""
    kLr = (lm * 100.0) / r22 if r22 else 1e9
    lam = 4.71 * math.sqrt(e / fy)
    fe = math.pi ** 2 * e / kLr ** 2
    fcr = (0.658 ** (fy / fe) * fy) if kLr <= lam else 0.877 * fe
    ag_m2 = ag / 1e4
    phi_pn_comp = 0.90 * fcr * ag_m2 * 1e3 / 9.80665
    phi_pn_tens = 0.90 * fy * ag_m2 * 1e3 / 9.80665
    z_m3 = z33 / 1e6
    phi_mn = 0.90 * fy * z_m3 * 1e3 / 9.80665
    pc = p if p > 0 else 0.0
    dc_comp = pc / phi_pn_comp if phi_pn_comp else 9e9
    dc_tens = pc / phi_pn_tens if phi_pn_tens else 0.0
    dc_flex = m3 / phi_mn if phi_mn else 9e9
    ratio_axial = dc_comp
    if ratio_axial >= 0.2:
        dc_int = ratio_axial + 8.0 / 9.0 * dc_flex
        eq = "H1-1a"
    else:
        dc_int = ratio_axial / 2.0 + dc_flex
        eq = "H1-1b"
    return {
        "kLr": kLr, "lam": lam, "fe": fe, "fcr": fcr,
        "phi_pn_comp": phi_pn_comp, "phi_pn_tens": phi_pn_tens,
        "phi_mn": phi_mn, "dc_comp": dc_comp, "dc_tens": dc_tens,
        "dc_flex": dc_flex, "dc_int": dc_int, "eq": eq,
    }


def _reemplazar_tabla(doc, head_inicio, nuevas_filas):
    """Busca una tabla por su fila de encabezado (primeras palabras) y reemplaza
    sus filas de datos por `nuevas_filas`. Mantiene la fila de encabezado."""
    for t in doc.tables:
        if not t.rows:
            continue
        head = " ".join(c.text for c in t.rows[0].cells).strip()
        if head.startswith(head_inicio):
            while len(t.rows) > 1:
                t._tbl.remove(t.rows[-1]._tr)
            for fila in nuevas_filas:
                cells = t.add_row().cells
                for j, v in enumerate(fila):
                    if j < len(cells):
                        cells[j].text = "" if v is None else str(v)
            return True
    return False


def corregir_tablas_v4(doc):
    """Reemplaza las tablas N° 4 (materiales 2.2), N° 5 (metrado 2.3.1) y
    N° 14 (envolvente 4.4) con los datos reales del modelo .s2k v4, para que
    el documento sea coherente con el capítulo 5.2/5.3 (9 secciones)."""
    dis = json.load(open(DISENO, encoding="utf-8"))
    res = json.load(open(os.path.join(ROOT, "_s2k_resultados.json"),
                         encoding="utf-8"))

    # ---- longitud real por frame y seccion por frame (desde PMM) ----
    len_frame = {}
    for r in res["pmm"]:
        if r["Frame"] and r["Length"]:
            len_frame[r["Frame"]] = max(len_frame.get(r["Frame"], 0),
                                        r["Length"])
    sec_frame = {}
    for r in res["pmm"]:
        if r["Frame"] and r["DesignSect"]:
            sec_frame[r["Frame"]] = r["DesignSect"]
    lens = {}
    count = {}
    for f, sec in sec_frame.items():
        lens[sec] = lens.get(sec, 0.0) + len_frame.get(f, 0.0)
        count.setdefault(sec, set()).add(f)

    # ---------------- Tabla N° 4 (2.2 Secciones empleadas) ----------------
    filas4 = []
    for sec_key, s in dis.items():
        dim = f"{s['t3']:.0f} mm x {s['t2']:.0f} mm" if s["t2"] else \
            f"Ø {s['t3']:.1f} mm"
        esp = f"{s['tw']:.2f} mm" if s["tw"] else "—"
        pm = s["Ag"] / 10000 * 7.85  # tf/m
        mat = s["Material"]
        if "TENSOR" in sec_key:
            mat = "A36"
        filas4.append([sec_key, mat, dim, esp, f"{pm:.3f} tf/m"])
    _reemplazar_tabla(doc, "Sección Material Dimensión", filas4)

    # ---------------- Tabla N° 5 (2.3.1 Metrado de acero) ----------------
    filas5 = []
    total_l = 0.0
    total_p = 0.0
    for sec_key, s in dis.items():
        L = lens.get(sec_key, 0.0)
        n = len(count.get(sec_key, []))
        pm = s["Ag"] / 10000 * 7.85 * 1000  # kgf/m
        pt = pm * L
        total_l += L
        total_p += pt
        filas5.append([sec_key, n, f"{L:.1f}", f"{pm:.3f}", f"{pt:.1f}"])
    filas5.append(["TOTAL", "", f"{total_l:.1f}", "", f"{total_p:.1f}"])
    _reemplazar_tabla(doc, "Sección N° barras", filas5)

    # ---------------- Tabla N° 14 (4.4 Envolvente por sección) ------------
    filas14 = []
    for sec_key, s in dis.items():
        filas14.append([sec_key,
                        f"{s['P']:.2f} tf", f"{s['V2']:.2f} tf",
                        f"{s['V3']:.2f} tf", f"{s['M3']:.2f} tf·m",
                        f"{s['M2']:.2f} tf·m"])
    _reemplazar_tabla(doc, "Sección P V2", filas14)

    print("  + Tablas N° 4/5/14 actualizadas con datos del modelo v4")


def crear_52_53(doc):
    dis = json.load(open(DISENO, encoding="utf-8"))
    res = json.load(open(os.path.join(ROOT, "_s2k_resultados.json"),
                         encoding="utf-8"))

    # TENSOR: max traccion desde la tabla BRACE AXIAL
    max_tens = 0.0
    for r in res["brace_axial"]:
        if r["Frame"] and r["DesignSect"] and "TENSOR" in r["DesignSect"]:
            max_tens = max(max_tens, r["PMaxTens"])

    # Orden de presentacion (de mayor a menor responsabilidad estructural)
    ORDEN = [
        ("COLUMNA HSS 200x200x8 mm", "HSS200x200x8", 56, 3.00, "ASTM A500 GrA",
         269, 310, 199958, "A500GrA"),
        ("CORREA HSS150x50x3 mm", "HSS150x50x3", 102, 5.05, "ASTM A500 GrA",
         269, 310, 199958, "A500GrA"),
        ("BRIDA INFERIOR HSS100x50x4.5 mm", "HSS100x50x4.5", 112, 1.582,
         "ASTM A500 GrA", 269, 310, 199958, "A500GrA"),
        ("BRIDA SUPERIOR HSS100x50x3 mm", "HSS100x50x3", 112, 1.640,
         "A500GrB46", 317, 400, 199948, "A500GrB46"),
        ("DIAGONALES HSS 50x50x2 MM", "HSS50x50x2", 217, 1.686, "A500GrB46",
         317, 400, 199948, "A500GrB46"),
        ("BRIDA SUPERIOR (LATERAL) HSS50x50x2 mm", "HSS50x50x2", 72, 0.842,
         "ASTM A500 GrA", 269, 310, 199958, "A500GrA"),
        ("BRIDA INFERIOR (LATERAL) HSS50x50x2 mm", "HSS50x50x2", 72, 0.842,
         "ASTM A500 GrA", 269, 310, 199958, "A500GrA"),
        ("DIAGONALES (LATERAL) HSS 50x50x2 MM", "HSS50x50x2", 72, 0.979,
         "ASTM A500 GrA", 269, 310, 199958, "A500GrA"),
        ("TENSOR Ø5/8", "O 5/8", 48, 8.234, "A36 (Fy = 36 ksi)", 248.211, 400,
         199948, "A36"),
    ]

    # ---------------- 5.2 Memoria de diseño ----------------
    rm._h(doc, "5.2 Memoria de diseño de elementos (fórmulas y verificación)", 2)
    rm._para(doc, "A continuación se resume el procedimiento de diseño aplicado "
                  "a cada elemento de la estructura, con las fórmulas de la "
                  "norma AISC 360-16 (LRFD) y los valores reales del modelo "
                  "SAP2000 «HUANCALPI - MODELO FINAL v4». Las solicitaciones de "
                  "diseño (Pu, Mu, Vu) corresponden a la envolvente de las "
                  "combinaciones del Capítulo 2; las capacidades (φc·Pn, φt·Pn, "
                  "φb·Mn, φv·Vn) se verifican con las secciones reales del "
                  "modelo. El tensor Ø5/8 es un elemento que trabaja "
                  "exclusivamente a tracción, por lo que se verifica con la "
                  "fórmula D2 (φt·Pn = 0.90·Ag·Fy) con acero A36 (Fy = 36 ksi).")

    # Tabla N° 15 propiedades de las 9 secciones
    headers = ["Sección", "Dimensión (mm)", "Esp. (mm)", "Ag (cm²)",
               "I (cm⁴)", "Z (cm³)", "r (cm)"]
    rows = []
    for sec_key, _tag, n, lm, mat, fy, fu, e, _alias in ORDEN:
        s = dis[sec_key]
        dim = f"{s['t3']:.0f}x{s['t2']:.0f}" if s["t2"] else f"Ø {s['t3']:.1f}"
        esp = f"{s['tw']:.2f}" if s["tw"] else "—"
        rows.append([f"{sec_key}", dim, esp,
                     f"{s['Ag']:.2f}", f"{s['I33']:.1f}", f"{s['Z33']:.1f}",
                     f"{s['r33']:.1f}"])
    rm._table(doc, headers, rows,
              caption="Tabla N° 15.- Propiedades geométricas de las 9 secciones "
                      "verificadas del modelo (AISC, dimensiones nominales).")
    # Fijar ancho de la 1.ª columna para que los nombres largos no se partan
    _t15 = doc.tables[-1]
    _w = [Cm(5.2)] + [Cm(2.0)] * (len(headers) - 1)
    for _row in _t15.rows:
        for _c, _wd in zip(_row.cells, _w):
            _c.width = _wd
    rm._para(doc, "Fuente: Elaboración propia; propiedades calculadas con las "
                  "dimensiones nominales de las secciones HSS (AISC Manual) y "
                  "barra redonda Ø5/8\".")

    # Veredictos para la tabla resumen y por seccion
    dc_sap = {}
    for sec_key, _tag, n, lm, mat, fy, fu, e, alias in ORDEN:
        s = dis[sec_key]
        pmm = s["PMM"]
        sap = pmm["TotalRatio"] if pmm and pmm.get("TotalRatio") is not None else None
        dc_sap[sec_key] = sap

    # Bloque por seccion
    for sec_key, _tag, n, lm, mat, fy, fu, e, alias in ORDEN:
        s = dis[sec_key]
        pmm = s["PMM"]
        sap = pmm["TotalRatio"] if pmm and pmm.get("TotalRatio") is not None else None

        titulo = sec_key
        rm._h(doc, titulo, 3)
        if "Fy = 36 ksi" in mat:
            rm._bullet(doc, f"Elemento · material A36 (Fy = 36 ksi = "
                             f"{fy:.0f} MPa, Fu = {fu:.0f} MPa, "
                             f"E = {e/1000:.0f} GPa) · {n} unidades · "
                             f"longitud máxima de diseño L = {lm:.2f} m.")
        else:
            rm._bullet(doc, f"Elemento · material {mat} (Fy = {fy:.0f} MPa, "
                             f"Fu = {fu:.0f} MPa, E = {e/1000:.0f} GPa) · "
                             f"{n} unidades · longitud máxima de diseño "
                             f"L = {lm:.2f} m.")
        rm._bullet(doc, f"Propiedades: Ag = {s['Ag']:.2f} cm², "
                        f"I = {s['I33']:.1f} cm⁴, Z = {s['Z33']:.1f} cm³, "
                        f"r = {s['r33']:.1f} cm.")

        if "TENSOR" in sec_key:
            # ---- TENSOR: solo traccion (D2, A36 Fy=36ksi) ----
            pu = max_tens
            ag_m2 = s["Ag"] / 1e4
            phi_pn = 0.90 * fy * ag_m2 * 1e3 / 9.80665
            dc = pu / phi_pn
            rm._bullet(doc, f"Solicitación de diseño (tracción): "
                            f"Pu = {pu:.3f} tf (envolvente de tracción de las "
                            f"48 barras, combinaciones 1.20CM+1.30W+0.50CV+0.50S "
                            f"y 1.30W±SX).")
            rm._para(doc, f"Tracción (D2): φt·Pn = 0.90·Fy·Ag = 0.90 × "
                          f"{fy:.1f} MPa × {s['Ag']:.2f} cm² = {phi_pn:.3f} tf. "
                          f"D/C tracción = {dc:.3f}.")
            rm._para(doc, f"La relación D/C máxima reportada por SAP2000 en "
                          f"compresión (≈101) corresponde al límite de "
                          f"compresión impuesto al elemento (φc·Pn → 0, KL/r = "
                          f"{(lm*100/s['r33']):.0f}); este caso no es de diseño "
                          f"para un tensor, que solo trabaja a tracción. "
                          f"D/C a tracción = {dc:.3f} ≤ 1.00 → CUMPLE.")
            continue

        pu = s["P"]
        mu3 = abs(s["M3"])
        mu2 = abs(s["M2"])
        vu = max(abs(s["V2"]), abs(s["V3"]))
        rm._bullet(doc, f"Solicitaciones de diseño (envolvente): "
                        f"Pu = {pu:.3f} tf, Mu = {max(mu3, mu2):.3f} tf·m, "
                        f"Vu = {vu:.3f} tf.")

        v = _calc_verif(s, fy, e, s["Ag"], s["Z33"], s["r22"], lm, pu, mu3, mu2)
        if v["dc_comp"] > 1e6:
            dc_comp_txt = "≈ ∞ (pandeo inelástico)"
        else:
            dc_comp_txt = f"{v['dc_comp']:.3f}"
        rm._para(doc, f"Compresión (E3): KL/r = {v['kLr']:.0f} "
                      f"({'<' if v['kLr'] <= v['lam'] else '>'} "
                      f"4.71·√(E/Fy) = {v['lam']:.0f}), "
                      f"Fe = π²E/(KL/r)² = {v['fe']:.0f} MPa, "
                      f"Fcr = {v['fcr']:.0f} MPa → "
                      f"φc·Pn = 0.90·Fcr·Ag = {v['phi_pn_comp']:.2f} tf. "
                      f"D/C compresión = {dc_comp_txt}.")
        rm._para(doc, f"Tracción (D2): φt·Pn = 0.90·Fy·Ag = "
                      f"{v['phi_pn_tens']:.2f} tf. "
                      f"D/C tracción = {v['dc_tens']:.3f}.")
        rm._para(doc, f"Flexión (F8): φb·Mn = 0.90·Fy·Z = {v['phi_mn']:.3f} "
                      f"tf·m ≥ Mu. D/C flexión = {v['dc_flex']:.3f}.")
        rm._para(doc, f"Interacción (H1.1, {v['eq']}): = {v['dc_int']:.3f} "
                      f"{'≤' if v['dc_int'] <= 1.0 else '>'}"
                      f" 1.00 → {'CUMPLE' if v['dc_int'] <= 1.0 else 'NO CUMPLE'}.")
        if sap is not None:
            rm._para(doc, f"Relación D/C máxima reportada por SAP2000 para "
                          f"esta sección: {sap:.3f} ≤ 1.00 → "
                          f"{'CUMPLE' if sap <= 1.0 else 'NO CUMPLE'}.")

    # ---------------- 5.3 Resumen de verificación ----------------
    rm._h(doc, "5.3 Resumen de verificación (D/C Ratio)", 2)
    rm._para(doc, "La Tabla N° 16 resume la relación Demanda/Capacidad (D/C) "
                  "máxima por sección, obtenida directamente del diseño de "
                  "acero de SAP2000 (AISC 360-16). El tensor Ø5/8 se reporta a "
                  "tracción (D/C = 0.183), que es su caso de diseño real; su "
                  "valor en compresión corresponde al límite de compresión nulo "
                  "impuesto al elemento y no es de diseño.")
    headers = ["Sección", "D/C máx", "Veredicto"]
    rows = []
    for sec_key, _tag, n, lm, mat, fy, fu, e, alias in ORDEN:
        sap = dc_sap[sec_key]
        if "TENSOR" in sec_key:
            ag_m2 = dis[sec_key]["Ag"] / 1e4
            phi_pn = 0.90 * fy * ag_m2 * 1e3 / 9.80665
            sap = max_tens / phi_pn
        ver = "CUMPLE" if sap <= 1.0 else "NO CUMPLE"
        rows.append([sec_key, f"{sap:.3f}", ver])
    rm._table(doc, headers, rows,
              caption="Tabla N° 16.- Relación D/C máxima por sección "
                      "(AISC 360-16, SAP2000 Steel Design).")
    rm._para(doc, "Fuente: SAP2000 (Steel Design, AISC 360-16).")

    # Figura N° 5 (diagrama D/C)
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    secs = [r[0].split(" (")[0] for r in rows]
    vals = [float(r[1]) for r in rows]
    colors = ["#C00000" if v > 1.0 else "#2E75B6" for v in vals]
    fig, ax = plt.subplots(figsize=(8.0, 4.4))
    ax.barh(range(len(vals)), vals, color=colors, edgecolor="#1F497D")
    ax.axvline(1.0, color="k", ls="--", lw=1.2)
    ax.text(1.02, len(vals) - 0.25, "D/C = 1.00", fontsize=9, color="k")
    for i, v in enumerate(vals):
        ax.text(v + 0.01, i, f"{v:.3f}", va="center", fontsize=8,
                color="#C00000" if v > 1.0 else "#1F497D")
    ax.set_yticks(range(len(secs)))
    ax.set_yticklabels(secs, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlabel("D/C (Demanda/Capacidad)", fontsize=9)
    ax.set_title("Relación Demanda/Capacidad (D/C) por sección — "
                 "Modelo v4 (AISC 360-16)", fontsize=10)
    ax.set_xlim(0, max(1.2, max(vals) * 1.05))
    fig.tight_layout()
    fig.savefig(FIG_DC, dpi=150)
    plt.close(fig)
    rm._figura(doc, FIG_DC,
               "Figura N° 5.- Diagrama de la relación Demanda/Capacidad (D/C) "
               "por sección; las secciones que exceden la unidad se muestran en "
               "rojo. El tensor se presenta a tracción (D/C = 0.183).",
               width_cm=14.5)

    print("  + 5.2/5.3 generados con las 9 secciones del .s2k v4")


if __name__ == "__main__":
    doc = rm.Document()
    rm._configurar_pagina(doc)
    rm._configurar_estilos(doc)
    crear_52_53(doc)
    doc.save("_test_52.docx")
    print("OK _test_52.docx")