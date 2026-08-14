"""Corrección DEFINITIVA - Renumeración completa + 6.4 expandido + imágenes.

ESTRUCTURA FINAL CORRECTA:
  1. METODOLOGÍA DE CÁLCULO ESTRUCTURAL
  2. GENERALIDADES
  3. ANÁLISIS POR CARGAS DE GRAVEDAD
  4. ANÁLISIS SÍSMICO Y DE VIENTO
  5. RESULTADOS DEL ANÁLISIS
  6. DISEÑO DE ELEMENTOS ESTRUCTURALES (6.4 CONEXIONES completo + 5 imágenes)
  7. CIMENTACIÓN
  8. CONCLUSIONES Y COMENTARIOS
  8. DRENAJE PLUVIAL
  9. CONTROL DE VERSIONES Y FIRMAS
  ANEXO A. PLANOS

FORMATO PRONIED:
  ✅ Arial en TODO
  ✅ Texto JUSTIFICADO
  ✅ Heading 1/2/3: Arial, #1F497D / #2E75B6
  ✅ Tablas "Tabla N° X" + fuente, encabezado #1F497D
  ✅ Figuras "Figura N° X" centradas, leyenda cursiva 9pt
  ✅ Texto JUSTIFICADO, interlineado 1.15, sangría 0.5cm
  ✅ 5 imágenes conexiones con rotulación "Figura N° X"
  ✅ Un solo TOC field (sin índice estático duplicado)
  ✅ Sin Capítulo 0 → empieza en 1. METODOLOGÍA
  ✅ Drenaje = Capítulo 8
"""

from __future__ import annotations

import copy
import os
import re
import shutil

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(ROOT, "Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.original.docx")
OUT = os.path.join(ROOT, "Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx")
CONN_PDF = os.path.join(ROOT, "Informes de conexión metálica columna PedestaL.pdf")
DRENAJE = os.path.join(ROOT, "informe_drenaje_pluvial.docx")
CONN_IMG_DIR = os.path.join(ROOT, "_conn_pdf_images")

S2K = os.path.join(ROOT, r"MN\HUANCALPI - MODELO FINAL v4.s2k")
CONN_PDF = os.path.join(ROOT, "Informes de conexión metálica columna PedestaL.pdf")
DRENAJE = os.path.join(ROOT, "informe_drenaje_pluvial.docx")
CONN_IMG_DIR = os.path.join(ROOT, "_conn_pdf_images")

_STATIC_INDICE = [
    "1.  GENERALES", "1.1  Alcance del documento", "1.2  Objetivos",
    "1.3  Descripción del sistema estructural",
    "1.4  Bases legales y normativa aplicable",
    "1.5  Información general del proyecto", "1.6  Materiales y perfiles",
    "1.7  Condiciones de cimentación",
    "2.  ANÁLISIS POR CARGAS DE GRAVEDAD", "2.1  Modelo estructural",
    "2.2  Secciones empleadas", "2.3  Metrado de cargas",
    "2.4  Combinaciones de carga",
    "3.  ANÁLISIS SÍSMICO Y DE VIENTO",
    "3.1  Parámetros sísmicos (NTP E.030)",
    "3.2  Periodo fundamental y masas participativas",
    "3.3  Cortante basal estático (verificación)",
    "3.4  Cargas de viento (NTP E.020)", "3.5  Control de desplazamientos",
    "4.  RESULTADOS DEL ANÁLISIS", "4.1  Periodos y modos de vibración",
    "4.2  Desplazamientos máximos", "4.3  Reacciones en apoyos",
    "4.4  Fuerzas internas envolventes por sección",
    "5.  DISEÑO DE ELEMENTOS ESTRUCTURALES", "5.1  Criterios de diseño",
    "5.2  Memoria de diseño de elementos (fórmulas y verificación)",
    "5.3  Resumen de verificación (D/C Ratio)",
    "5.4  Conexiones, placas base y pernos de anclaje",
    "5.5  Cimentación",
    "6.  CONCLUSIONES Y COMENTARIOS", "ANEXO A  PLANOS",
]

AZUL_N1 = RGBColor(0x1F, 0x49, 0x7D)
AZUL_N2 = RGBColor(0x2E, 0x75, 0xB6)
FONDO_TABLA = "1F497D"
FONDO_CELDA = "D9E2F3"


def _configurar_estilos(doc: Document):
    for style in doc.styles:
        if style.type == 1:
            style.font.name = 'Arial'
    
    s1 = doc.styles['Heading 1']
    s1.font.name = 'Arial'
    s1.font.size = Pt(14)
    s1.font.bold = True
    s1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    s1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s1.paragraph_format.space_before = Pt(18)
    s1.paragraph_format.space_after = Pt(6)
    s1.paragraph_format.line_spacing = Pt(16)
    
    s2 = doc.styles['Heading 2']
    s2.font.name = 'Arial'
    s2.font.size = Pt(12)
    s2.font.bold = True
    s2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    s2.paragraph_format.space_before = Pt(12)
    s2.paragraph_format.space_after = Pt(4)
    
    s3 = doc.styles['Heading 3']
    s3.font.name = 'Arial'
    s3.font.size = Pt(11)
    s3.font.bold = True
    s3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = Pt(12)
    normal.paragraph_format.line_spacing_rule = 1
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.5)
    
    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Arial'
        lb.font.size = Pt(10)
        lb.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        lb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lb.paragraph_format.line_spacing = Pt(12)
        lb.paragraph_format.left_indent = Cm(1.0)
        lb.paragraph_format.first_line_indent = Cm(-0.5)


def _h(doc, text, level):
    return doc.add_paragraph(text, style=f"Heading {level}")


def _para(doc, text, size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=Pt(0), space_after=Pt(4), italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.line_spacing_rule = 1
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = size
    run.bold = bold
    run.italic = False
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def _bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def _table(doc, headers, rows, caption=None, size=Pt(8.5)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "1F497D")
        shading.set(qn('w:val'), 'clear')
        hdr[j]._tc.get_or_add_tcPr().append(shading)
    
    for i, r in enumerate(rows):
        row = t.add_row()
        for j, v in enumerate(r):
            row.cells[j].text = "" if v is None else str(v)
            for par in row.cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D9E2F3")
                shading.set(qn('w:val'), 'clear')
                row.cells[j]._tc.get_or_add_tcPr().append(shading)
    
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return t


def _insertar_imagen(doc, img_path, caption, width_cm=14):
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(14))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(caption)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _limpiar_indice_estatico(doc):
    for p in list(doc.paragraphs):
        if p.text.strip() in _STATIC_INDICE:
            p._p.getparent().remove(p._p)
    return True


def _renumerar_capitulos(doc):
    """Renumera todos los encabezados Heading 1, 2 y 3 sumando 1 al número de capítulo."""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            m = re.match(r'^(\d+)\.\s+(.+)$', p.text.strip())
            if m:
                num = int(m.group(1)) + 1
                p.text = f"{num}. {m.group(2)}"
                # Mantener formato
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif p.style.name == "Heading 2":
            m = re.match(r'^(\d+)\.(\d+)\s+(.+)$', p.text.strip())
            if m:
                cap = int(m.group(1)) + 1
                sub = m.group(2)
                texto = m.group(3)
                p.text = f"{cap}.{sub} {texto}"
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif p.style.name == "Heading 3":
            m = re.match(r'^(\d+)\.(\d+)\.(\d+)\s+(.+)$', p.text.strip())
            if m:
                cap = int(m.group(1)) + 1
                sub = m.group(2)
                subsub = m.group(3)
                texto = m.group(4)
                p.text = f"{cap}.{sub}.{subsub} {texto}"
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def _h(doc, text, level):
    return doc.add_paragraph(text, style=f"Heading {level}")


def _para(doc, text, size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=Pt(0), space_after=Pt(4), italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.line_spacing_rule = 1
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = size
    run.bold = bold
    run.italic = False
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def _bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def _table(doc, headers, rows, caption=None, size=Pt(8.5)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "1F497D")
        shading.set(qn('w:val'), 'clear')
        hdr[j]._tc.get_or_add_tcPr().append(shading)
    
    for i, r in enumerate(rows):
        row = t.add_row()
        for j, v in enumerate(r):
            row.cells[j].text = "" if v is None else str(v)
            for par in row.cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D9E2F3")
                shading.set(qn('w:val'), 'clear')
                row.cells[j]._tc.get_or_add_tcPr().append(shading)
    
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return t


def _insertar_imagen(doc, img_path, caption, width_cm=14):
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(14))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(caption)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _limpiar_indice_estatico(doc):
    for p in list(doc.paragraphs):
        if p.text.strip() in _STATIC_INDICE:
            p._p.getparent().remove(p._p)
    return True


def _h(doc, text, level):
    return doc.add_paragraph(text, style=f"Heading {level}")


def _para(doc, text, size=Pt(10), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=Pt(0), space_after=Pt(4), italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.line_spacing_rule = 1
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = size
    run.bold = bold
    run.italic = False
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def _bullet(doc, text, size=Pt(10)):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def _table(doc, headers, rows, caption=None, size=Pt(8.5)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        for par in hdr[j].paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "1F497D")
        shading.set(qn('w:val'), 'clear')
        hdr[j]._tc.get_or_add_tcPr().append(shading)
    
    for i, r in enumerate(rows):
        row = t.add_row()
        for j, v in enumerate(r):
            row.cells[j].text = "" if v is None else str(v)
            for par in row.cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if i % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D9E2F3")
                shading.set(qn('w:val'), 'clear')
                row.cells[j]._tc.get_or_add_tcPr().append(shading)
    
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return t


def _insertar_imagen(doc, img_path, caption, width_cm=14):
    if not os.path.exists(img_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(14))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(caption)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _limpiar_indice_estatico(doc):
    for p in list(doc.paragraphs):
        if p.text.strip() in _STATIC_INDICE:
            p._p.getparent().remove(p._p)
    return True


# ----------------------------------------------------- CONTENIDO A INSERTAR

def _capitulo_1_metodologia(doc):
    _h(doc, "1. METODOLOGÍA DE CÁLCULO ESTRUCTURAL", 1)
    _para(doc, "Esta sección reúne la metodología de cálculo estructural de la "
               "cobertura metálica en arco-tijeral, de carácter transversal a los "
               "capítulos siguientes. La memoria se genera de forma automatizada "
               "a partir del modelo SAP2000 y se complementa con el informe de "
               "conexiones (IDEA StatiCa) y el de drenaje pluvial.")
    _h(doc, "1.1 Procedimiento de generación", 2)
    _bullet(doc, "Flujo: .s2k (modelo+resultados) → sap2000gen/memoria/ → "
                  ".docx (formato PRONIED).")
    _bullet(doc, "Comando: `python cli.py memoria --model "
                  "'MN\\HUANCALPI - MODELO FINAL v4.s2k' -o "
                  "'Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx' "
                  "--formato pronied`.")
    _bullet(doc, "Lectura de PDFs: PyMuPDF extrae texto; si la página es "
                  "escaneada, la convierte a PNG → `cli.py pdf`.")
    _h(doc, "1.2 Unidades", 2)
    _bullet(doc, "Modelo SAP2000: Tonf, m, C.")
    _bullet(doc, "Memoria: tf·m, MPa, kN, cm, kgf. Cap. 6 conserva unidades "
                  "nativas del reporte IDEA StatiCa (ksi, kN, cm); 1 ksi ≈ "
                  "6.895 MPa. Cap. 8 (Drenaje): m², mm/h, L/s, m.")
    _h(doc, "1.3 Metodología estructural", 2)
    _bullet(doc, f"Origen: '{S2K}' (unidades Tonf, m, C).")
    _bullet(doc, "Elementos: 414 nudos, 863 barras, 96 áreas, 14 apoyos; "
                  "acero 13.688 tf.")
    _bullet(doc, "Geometría (m): 22.65 × 30.30 × 11.86 (Z 0.20→11.86).")
    _bullet(doc, "Diseño: SAP2000 v27 + AISC 360-16 (LRFD); verificación "
                  "D/C ≤ 1.00 (φ=0.90 C/B/T/V).")
    _bullet(doc, "22 combinaciones LRFD: 1.4CM; 1.2CM+1.6CV+0.5S; "
                  "1.2CM+1.3W+0.5CV+0.5S; 1.2CM+1.6S+(0.5CV ó 0.80W); "
                  "envolventes 1.3W ó +SX.")
    _bullet(doc, "Normas: NTP E.020/030/050/060/090; AISC 341-16,303; "
                  "ASCE 7-16; AWS D1.1; ASTM A500/A572/F1554/A325.")
    _bullet(doc, "Cimentación: zapatas 1.50×1.50 m @ 1.50 m; qadm=0.50 "
                  "kg/cm² (FS=3.0); asent. 0.802 cm < 2.50 cm.")
    _bullet(doc, "Desplazamientos: derivas NTP E.030; flecha L/200 (gravedad) "
                  "/ L/250 (viento); T₁=0.625 s; masa ≥90 %.")


def _crear_conexiones_completo(doc):
    _para(doc, "La placa base de las columnas se verificó con el modelo de "
               "elementos finitos 'Coneccion plancha base.ideaCon' (IDEA "
               "StatiCa, método CBFEM); el informe completo se adjunta como "
               f"archivo '{os.path.basename(CONN_PDF)}'. Las conexiones se "
               "diseñan soldadas (AWS D1.1) o empernadas (ASTM A325).")
    _h(doc, "6.4.1 Datos de diseño y materiales", 3)
    _table(doc, ["Concepto", "Valor", "Unidad"], [
        ["Proyecto", "HUANCALPI-RIBBECK", "—"],
        ["Norma estructural", "AISC 360-16 (LRFD)", "—"],
        ["Norma anclaje", "ACI 318-14 §17.4", "—"],
        ["Sección transversal columna", "HSS(Imp) 8×8×5/8", "—"],
        ["Material columna", "A36", "Fy = 36 ksi (εlim = 5 %)"],
        ["Perno anclaje", "¾\" A325", "fu = 119.7 ksi; Área = 2.850 cm²"],
        ["Longitud anclaje", "100.0", "cm"],
        ["Transferencia a corte", "Fricción / junta de mortero (2.40 cm)", "—"],
        ["Hormigón de apoyo", "4000 / 3000 psi", "—"],
    ], caption="Tabla 6.4.1 — Datos de diseño y materiales de la conexión.")

    _h(doc, "6.4.2 Cargas de equilibrio (LE1) — Joint COL", 3)
    _table(doc, ["Fuerza", "Valor", "Unidad"], [
        ["N (compresión)", "-66,539", "kN"],
        ["Vy", "4.901", "kN"],
        ["Vz", "-29.184", "kN"],
        ["Mx", "0.00", "kN·m"],
        ["My", "-14.03", "kN·m"],
        ["Mz", "-85.66", "kN·m"],
    ], caption="Tabla 6.4.2 — Fuerzas de equilibrio en la columna.")

    _h(doc, "6.4.3 Resumen de verificaciones", 3)
    _table(doc, ["Concepto", "Ratio", "Límite", "Estado"], [
        ["Análisis general", "100.0", "%", "OK"],
        ["Placas (BP1)", "0.7", "< 5.0 %", "OK"],
        ["Anclajes (tracción)", "99.3", "< 100 %", "OK"],
        ["Soldaduras", "88.4", "< 100 %", "OK"],
        ["Bloque de hormigón", "12.9", "< 100 %", "OK"],
        ["Corte bloque-hormigón", "36.7", "< 100 %", "OK"],
        ["Pandeo", "—", "No calculado", "—"],
    ], caption="Tabla 6.4.3 — Resumen de estados de verificación.")

    _h(doc, "6.4.4 Verificación por elemento", 3)
    _table(doc, ["Elemento", "Fy (ksi)", "Espesor", "σEd (ksi)",
                 "εPl (%)", "σcEd (ksi)", "Estado"], [
        ["COL", "36.0", "5/8\"", "32.4", "0.0", "0.0", "OK"],
        ["BP1 (placa base)", "36.0", "1\"", "28.5", "0.0", "0.0", "OK"],
        ["RIB2a", "36.0", "1/4\"", "32.4", "0.1", "0.0", "OK"],
        ["RIB2b", "36.0", "1/4\"", "32.5", "0.3", "0.0", "OK"],
        ["RIB4a", "36.0", "1/4\"", "32.5", "0.2", "0.0", "OK"],
        ["RIB4b", "36.0", "1/4\"", "32.6", "0.7", "0.0", "OK"],
    ], caption="Tabla 6.4.4 — Tensión/compresión por elemento (LE1).")

    _h(doc, "6.4.5 Verificación de anclajes", 3)
    _table(doc, ["Concepto", "ϕRn (kN)", "≥ Demanda (kN)", "Norma", "Estado"], [
        ["Tracción anclaje (A2)", "124,440", "84,262",
         "ACI 318-14 §17.4.1 (ϕ=0.70)", "OK"],
        ["Arrancamiento tracción grupo (A2,A3,A7)", "239,042", "237,345",
         "ACI 318-14 §17.4.2", "OK"],
        ["Arrancamiento lateral grupo", "984,889", "84,262",
         "ACI 318-14 §17.4.4 (ψ=0.51)", "OK"],
        ["Arrancamiento barra (A2-A7)", "1,125,290", "0.0",
         "ACI 318-14 §17.4.3", "OK"],
        ["Pryoc/hormigón (grupo)", "580,139", "0.0",
         "ACI 318-14 §17.5.3", "OK"],
        ["Aplastamiento bloque-hormigón", "5.1 ksi", "0.4 ksi",
         "AISC J8", "OK"],
        ["Interacción tracción-cortante (ACI 17.6)", "0.99", "≤ 1.0",
         "ACI 318-14 §R17.6", "OK"],
    ], caption="Tabla 6.4.5 — Verificación de anclajes a tracción/lateral.")
    _para(doc, "Fórmula anclaje tracción (ACI 318-14 §17.4.1): "
               "ϕNsa = ϕ·Asa·fu = 0.70 × 2.155 cm² × 119.7 ksi. "
               "Cono de arranque (§17.4.2): Acon = 12 636,5 cm²; "
               "Aco (individual) = 9 409 cm²; Ψed,N = 0.98. "
               "Lateral (§17.4.4): Área portante cabeza = 97 150 cm²; "
               "ψ = 0.51 (c_a1=46 cm, c_a2=48.5 cm, s=15 cm). "
               "Pryoc (§17.5.3): ϕ = 0.65; k_cp = 2.0; ϕVcp = 580 139 kN.",
               size=Pt(8.5))

    _h(doc, "6.4.6 Verificación de soldaduras", 3)
    _table(doc, ["Conexión", "ϕRn (kN)", "≥ Demanda (kN)", "Áng. (°)", "Estado"], [
        ["BP1 / RIB2a", "11,392", "8,649", "83.8", "OK"],
        ["COL-w 2 / RIB2a", "10,008", "7,630", "46.5", "OK"],
        ["BP1 / RIB2b", "11,358", "8,927", "81.1", "OK"],
        ["COL-w 2 / RIB2b", "11,392", "8,649", "83.8", "OK"],
        ["BP1 / RIB4a", "11,369", "8,913", "81.9", "OK"],
        ["COL-w 4 / RIB4a", "10,050", "7,689", "47.2", "OK"],
        ["COL-w 4 / RIB4b", "10,089", "7,835", "47.8", "OK"],
        ["BP1 / COL", "6,687", "5,910", "79.1", "OK"],
    ], caption="Tabla 6.4.6 — Resistencia de soldaduras (ϕ = 0.75, E70xx "
               "70 ksi) según AISC 360-16 J2-4 / AWS D1.1.")

    _h(doc, "6.4.7 Verificación del bloque de cimentación", 3)
    _table(doc, ["Concepto", "ϕRn", "≥ Demanda", "Norma", "Estado"], [
        ["Compresión bloque (BP1)", "5.1 ksi", "0.4 ksi", "AISC J8", "OK"],
        ["Corte bloque-hormigón", "—", "—", "ACI 318-14 (ϕVcp)", "OK"],
    ], caption="Tabla 6.4.7 — Verificación del bloque de cimentación.")
    _para(doc, "Bloque: 127×127 cm; profundidad 150 cm; área contacto "
               "1049,627 cm²; superficie apoyo 12 734,202 cm²; ϕ = 0.65; "
               "transferencia corte por fricción / junta de mortero (2.40 cm).",
               size=Pt(8.5))


def _crear_drenaje(doc):
    _h(doc, "8. DRENAJE PLUVIAL", 1)
    _para(doc, "Documento complementario de drenaje pluvial unificado a la "
               "memoria de cálculo. Unidades: m², mm/h, L/s, m.")
    _h(doc, "8.1 Resumen ejecutivo", 2)
    _bullet(doc, "Normativa: CE.040 Drenaje Pluvial (RNE RM 126-2021).")
    _bullet(doc, "Método racional: Q = 0.95·C·I·A/3600 (L/s); C = 0.95.")
    _bullet(doc, "Intensidad IILA-SENAMHI-UNI: I ≈ 30.33 mm/h (T = 25 a, "
                  "tc = 10 min, zona I23₁₀, Huancavelica).")
    _bullet(doc, "Área captación: A = 80,295 m² (proyección horizontal).")
    _bullet(doc, "Q_total = 0.643 L/s → 0.161 L/s/c/bajante (4 unidades).")
    _bullet(doc, "Canaleta caja 0.25×0.14 m (1 %), OK por Manning.")
    _bullet(doc, "Bajantes Ø50 mm PVC (n = 0.010): 2 por canalón "
                  "(criterio constructivo).")
    if os.path.exists(DRENAJE):
        _h(doc, "8.2 Texto íntegro del informe", 2)
        from sap2000gen.text_readers import docx_to_text
        # Filtrar: saltar el índice interno (páginas 15-34 del PDF original)
        # Criterio: saltar párrafos que sean solo "ÍNDICE" o entradas de tabla de contenido
        for num, t in docx_to_text(DRENAJE):
            txt = t.strip()
            # Saltar el título "ÍNDICE" solo
            if txt == "ÍNDICE" or txt == "INDICE":
                continue
            # Saltar entradas de tabla de contenido (líneas que empiezan con número + tab + texto + número de página)
            if re.match(r'^\d+\.\s+.+\t\d+$', txt) or re.match(r'^\d+\.\d+\.\s+.+\t\d+$', txt):
                continue
            if t.strip():
                _para(doc, t, size=Pt(9))
    else:
        _bullet(doc, f"No se encontró {DRENAJE}.")


# ----------------------------------------------------- BUILD PRINCIPAL

def build():
    # 1. Copia limpia del original
    if os.path.exists(OUT):
        os.remove(OUT)
    shutil.copy2(SRC, OUT)
    
    doc = Document(OUT)
    
    # 1. Configurar estilos globales
    _configurar_estilos(doc)
    
    # 1. ELIMINAR índice estático duplicado
    n = 0
    for p in list(doc.paragraphs):
        if p.text.strip() in _STATIC_INDICE:
            p._p.getparent().remove(p._p)
    print(f"  - Índice estático suprimido: {len(_STATIC_INDICE)} renglones")

    # 2. RENUMERAR todos los capítulos (+1) porque insertaremos Capítulo 1 al inicio
    _renumerar_capitulos(doc)
    print("  [OK] Capítulos renumerados (+1)")

    # 2. Insertar Capítulo 1 METODOLOGÍA antes de "2. GENERALES" (que era "1. GENERALES")
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip() == "2. GENERALES":
            anchor_idx = i
            break
    
    if anchor_idx is not None:
        aux = Document()
        _h(aux, "1. METODOLOGÍA DE CÁLCULO ESTRUCTURAL", 1)
        _para(aux, "Esta sección reúne la metodología de cálculo estructural de la "
                   "cobertura metálica en arco-tijeral, de carácter transversal a los "
                   "capítulos siguientes. La memoria se genera de forma automatizada "
                   "a partir del modelo SAP2000 y se complementa con el informe de "
                   "conexiones (IDEA StatiCa) y el de drenaje pluvial.")
        _h(aux, "1.1 Procedimiento de generación", 2)
        _bullet(aux, "Flujo: .s2k (modelo+resultados) → sap2000gen/memoria/ → "
                      ".docx (formato PRONIED).")
        _bullet(aux, "Comando: `python cli.py memoria --model "
                      "'MN\\HUANCALPI - MODELO FINAL v4.s2k' -o "
                      "'Memoria_de_Calculo_Estructural_Cobertura_Huancalpi.docx' "
                      "--formato pronied`.")
        _bullet(aux, "Lectura de PDFs: PyMuPDF extrae texto; si la página es "
                      "escaneada, la convierte a PNG → `cli.py pdf`.")
        _h(aux, "1.2 Unidades", 2)
        _bullet(aux, "Modelo SAP2000: Tonf, m, C.")
        _bullet(aux, "Memoria: tf·m, MPa, kN, cm, kgf. Cap. 6 conserva unidades "
                      "nativas del reporte IDEA StatiCa (ksi, kN, cm); 1 ksi ≈ "
                      "6.895 MPa. Cap. 8 (Drenaje): m², mm/h, L/s, m.")
        _h(aux, "1.3 Metodología estructural", 2)
        _bullet(aux, f"Origen: '{S2K}' (unidades Tonf, m, C).")
        _bullet(aux, "Elementos: 414 nudos, 863 barras, 96 áreas, 14 apoyos; "
                      "acero 13.688 tf.")
        _bullet(aux, "Geometría (m): 22.65 × 30.30 × 11.86 (Z 0.20→11.86).")
        _bullet(aux, "Diseño: SAP2000 v27 + AISC 360-16 (LRFD); verificación "
                      "D/C ≤ 1.00 (φ=0.90 C/B/T/V).")
        _bullet(aux, "22 combinaciones LRFD: 1.4CM; 1.2CM+1.6CV+0.5S; "
                      "1.2CM+1.3W+0.5CV+0.5S; 1.2CM+1.6S+(0.5CV ó 0.80W); "
                      "envolventes 1.3W ó +SX.")
        _bullet(aux, "Normas: NTP E.020/030/050/060/090; AISC 341-16,303; "
                      "ASCE 7-16; AWS D1.1; ASTM A500/A572/F1554/A325.")
        _bullet(aux, "Cimentación: zapatas 1.50×1.50 m @ 1.50 m; qadm=0.50 "
                      "kg/cm² (FS=3.0); asent. 0.802 cm < 2.50 cm.")
        _bullet(aux, "Desplazamientos: derivas NTP E.030; flecha L/200 (gravedad) "
                      "/ L/250 (viento); T₁=0.625 s; masa ≥90 %.")
        
        anchor = doc.paragraphs[anchor_idx]._p
        nodos = [copy.deepcopy(n) for n in aux._element.body]
        anchor.addprevious(nodos[0])
        for k in range(1, len(nodos)):
            nodos[k - 1].addnext(nodos[k])
        print("  + Capítulo 1 (Metodología) insertado antes de '2. GENERALIDADES'")

    doc.save(OUT)
    
    # RELOAD
    doc = Document(OUT)
    
    # 2. REEMPLAZAR 6.4 CONEXIONES (buscar "6.4 Conexiones, placas base..." que era "5.4")
    # Buscar el heading "6.4 Conexiones, placas base y pernos de anclaje" (Heading 2)
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 2" and "Conexiones, placas base" in p.text:
            anchor_idx = i
            break
    
    if anchor_idx is not None:
        # Eliminar el párrafo original "6.4 Conexiones..." y el siguiente párrafo (el texto corto)
        # e insertar el contenido completo
        p_original = doc.paragraphs[anchor_idx]
        p_next = doc.paragraphs[anchor_idx + 1] if anchor_idx + 1 < len(doc.paragraphs) else None
        
        # Eliminar párrafo original y el siguiente si es texto descriptivo
        p_original._p.getparent().remove(p_original._p)
        if p_next and p_next.style.name == "Normal" and len(p_next.text) < 200:
            p_next._p.getparent().remove(p_next._p)
        
        # Insertar contenido completo
        aux = Document()
        _para(aux, "La placa base de las columnas se verificó con el modelo de "
                   "elementos finitos 'Coneccion plancha base.ideaCon' (IDEA "
                   "StatiCa, método CBFEM); el informe completo se adjunta como "
                   f"archivo '{os.path.basename(CONN_PDF)}'. Las conexiones se "
                   "diseñan soldadas (AWS D1.1) o empernadas (ASTM A325).")
        _h(aux, "6.4.1 Datos de diseño y materiales", 3)
        _table(aux, ["Concepto", "Valor", "Unidad"], [
            ["Proyecto", "HUANCALPI-RIBBECK", "—"],
            ["Norma estructural", "AISC 360-16 (LRFD)", "—"],
            ["Norma anclaje", "ACI 318-14 §17.4", "—"],
            ["Sección transversal columna", "HSS(Imp) 8×8×5/8", "—"],
            ["Material columna", "A36", "Fy = 36 ksi (εlim = 5 %)"],
            ["Perno anclaje", "¾\" A325", "fu = 119.7 ksi; Área = 2.850 cm²"],
            ["Longitud anclaje", "100.0", "cm"],
            ["Transferencia a corte", "Fricción / junta de mortero (2.40 cm)", "—"],
            ["Hormigón de apoyo", "4000 / 3000 psi", "—"],
        ], caption="Tabla 6.4.1 — Datos de diseño y materiales de la conexión.")

        _h(aux, "6.4.2 Cargas de equilibrio (LE1) — Joint COL", 3)
        _table(aux, ["Fuerza", "Valor", "Unidad"], [
            ["N (compresión)", "-66,539", "kN"],
            ["Vy", "4.901", "kN"],
            ["Vz", "-29.184", "kN"],
            ["Mx", "0.00", "kN·m"],
            ["My", "-14.03", "kN·m"],
            ["Mz", "-85.66", "kN·m"],
        ], caption="Tabla 6.4.2 — Fuerzas de equilibrio en la columna.")

        _h(aux, "6.4.3 Resumen de verificaciones", 3)
        _table(aux, ["Concepto", "Ratio", "Límite", "Estado"], [
            ["Análisis general", "100.0", "%", "OK"],
            ["Placas (BP1)", "0.7", "< 5.0 %", "OK"],
            ["Anclajes (tracción)", "99.3", "< 100 %", "OK"],
            ["Soldaduras", "88.4", "< 100 %", "OK"],
            ["Bloque de hormigón", "12.9", "< 100 %", "OK"],
            ["Corte bloque-hormigón", "36.7", "< 100 %", "OK"],
            ["Pandeo", "—", "No calculado", "—"],
        ], caption="Tabla 6.4.3 — Resumen de estados de verificación.")

        _h(aux, "6.4.4 Verificación por elemento", 3)
        _table(aux, ["Elemento", "Fy (ksi)", "Espesor", "σEd (ksi)",
                     "εPl (%)", "σcEd (ksi)", "Estado"], [
            ["COL", "36.0", "5/8\"", "32.4", "0.0", "0.0", "OK"],
            ["BP1 (placa base)", "36.0", "1\"", "28.5", "0.0", "0.0", "OK"],
            ["RIB2a", "36.0", "1/4\"", "32.4", "0.1", "0.0", "OK"],
            ["RIB2b", "36.0", "1/4\"", "32.5", "0.3", "0.0", "OK"],
            ["RIB4a", "36.0", "1/4\"", "32.5", "0.2", "0.0", "OK"],
            ["RIB4b", "36.0", "1/4\"", "32.6", "0.7", "0.0", "OK"],
        ], caption="Tabla 6.4.4 — Tensión/compresión por elemento (LE1).")

        _h(aux, "6.4.5 Verificación de anclajes", 3)
        _table(aux, ["Concepto", "ϕRn (kN)", "≥ Demanda (kN)", "Norma", "Estado"], [
            ["Tracción anclaje (A2)", "124,440", "84,262",
             "ACI 318-14 §17.4.1 (ϕ=0.70)", "OK"],
            ["Arrancamiento tracción grupo (A2,A3,A7)", "239,042", "237,345",
             "ACI 318-14 §17.4.2", "OK"],
            ["Arrancamiento lateral grupo", "984,889", "84,262",
             "ACI 318-14 §17.4.4 (ψ=0.51)", "OK"],
            ["Arrancamiento barra (A2-A7)", "1,125,290", "0.0",
             "ACI 318-14 §17.4.3", "OK"],
            ["Pryoc/hormigón (grupo)", "580,139", "0.0",
             "ACI 318-14 §17.5.3", "OK"],
            ["Aplastamiento bloque-hormigón", "5.1 ksi", "0.4 ksi",
             "AISC J8", "OK"],
            ["Interacción tracción-cortante (ACI 17.6)", "0.99", "≤ 1.0",
             "ACI 318-14 §R17.6", "OK"],
        ], caption="Tabla 6.4.5 — Verificación de anclajes a tracción/lateral.")
        _para(aux, "Fórmula anclaje tracción (ACI 318-14 §17.4.1): "
                   "ϕNsa = ϕ·Asa·fu = 0.70 × 2.155 cm² × 119.7 ksi. "
                   "Cono de arranque (§17.4.2): Acon = 12 636,5 cm²; "
                   "Aco (individual) = 9 409 cm²; Ψed,N = 0.98. "
                   "Lateral (§17.4.4): Área portante cabeza = 97 150 cm²; "
                   "ψ = 0.51 (c_a1=46 cm, c_a2=48.5 cm, s=15 cm). "
                   "Pryoc (§17.5.3): ϕ = 0.65; k_cp = 2.0; ϕVcp = 580 139 kN.",
                   size=Pt(8.5))

        _h(aux, "6.4.6 Verificación de soldaduras", 3)
        _table(aux, ["Conexión", "ϕRn (kN)", "≥ Demanda (kN)", "Áng. (°)", "Estado"], [
            ["BP1 / RIB2a", "11,392", "8,649", "83.8", "OK"],
            ["COL-w 2 / RIB2a", "10,008", "7,630", "46.5", "OK"],
            ["BP1 / RIB2b", "11,358", "8,927", "81.1", "OK"],
            ["COL-w 2 / RIB2b", "11,392", "8,649", "83.8", "OK"],
            ["BP1 / RIB4a", "11,369", "8,913", "81.9", "OK"],
            ["COL-w 4 / RIB4a", "10,050", "7,689", "47.2", "OK"],
            ["COL-w 4 / RIB4b", "10,089", "7,835", "47.8", "OK"],
            ["BP1 / COL", "6,687", "5,910", "79.1", "OK"],
        ], caption="Tabla 6.4.6 — Resistencia de soldaduras (ϕ = 0.75, E70xx "
                   "70 ksi) según AISC 360-16 J2-4 / AWS D1.1.")

        _h(aux, "6.4.7 Verificación del bloque de cimentación", 3)
        _table(aux, ["Concepto", "ϕRn", "≥ Demanda", "Norma", "Estado"], [
            ["Compresión bloque (BP1)", "5.1 ksi", "0.4 ksi", "AISC J8", "OK"],
            ["Corte bloque-hormigón", "—", "—", "ACI 318-14 (ϕVcp)", "OK"],
        ], caption="Tabla 6.4.7 — Verificación del bloque de cimentación.")
        _para(aux, "Bloque: 127×127 cm; profundidad 150 cm; área contacto "
                   "1049,627 cm²; superficie apoyo 12 734,202 cm²; ϕ = 0.65; "
                   "transferencia corte por fricción / junta de mortero (2.40 cm).",
                   size=Pt(8.5))
        
        anchor = doc.paragraphs[anchor_idx]._p
        nodos = [copy.deepcopy(n) for n in aux._element.body]
        anchor.addprevious(nodos[0])
        for k in range(1, len(nodos)):
            nodos[k - 1].addnext(nodos[k])
        print("  + Capítulo 6.4 (conexiones) REEMPLAZADO con 7 subsecciones y 7 tablas")

    doc.save(OUT)
    
    # RELOAD para insertar imágenes en 6.4
    doc = Document(OUT)
    
    # Insertar imágenes en 6.4 después de cada subsección
    img_map = [
        ("conn_p17_img0.png", "Figura 6.4.1 — Modelo CBFEM de la conexión (página 17 del reporte IDEA StatiCa).", "6.4.3 Resumen de verificaciones"),
        ("conn_p7_img0.png", "Figura 6.4.2 — Verificación de anclajes a tracción (ACI 318-14 §17.4.1).", "6.4.5 Verificación de anclajes"),
        ("conn_p17_img1.png", "Figura 6.4.3 — Verificación de soldaduras fillet (AISC 360-16 J2-4).", "6.4.6 Verificación de soldaduras"),
        ("conn_p22_img0.png", "Figura 6.4.4 — Verificación del bloque de hormigón (AISC J8).", "6.4.7 Verificación del bloque"),
        ("conn_p28_img0.png", "Figura 6.4.5 — Modelo CBFEM: vistas 3D y detalle de soldaduras/anclajes (teoría).", "6.4.7 Verificación del bloque"),
    ]
    
    for img_name, caption, anchor_text in img_map:
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == "Heading 3" and anchor_text in p.text:
                img_path = os.path.join(CONN_IMG_DIR, img_name)
                if os.path.exists(img_path):
                    # Insertar imagen y caption DESPUÉS del heading
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(4)
                    p_img.paragraph_format.space_after = Pt(2)
                    run = p_img.add_run()
                    run.add_picture(os.path.join(CONN_IMG_DIR, img_name), width=Cm(14))
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(0)
                    p_cap.paragraph_format.space_after = Pt(8)
                    p_cap.paragraph_format.line_spacing = Pt(11)
                    run = p_cap.add_run(caption)
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    # Insertar imagen y caption después del heading
                    img_path = os.path.join(CONN_IMG_DIR, img_name)
                    if os.path.exists(img_path):
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(4)
                        p_img.paragraph_format.space_after = Pt(2)
                        run = p_img.add_run()
                        run.add_picture(os.path.join(CONN_IMG_DIR, img_name), width=Cm(14))
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(0)
                        p_cap.paragraph_format.space_after = Pt(8)
                        p_cap.paragraph_format.line_spacing = Pt(11)
                        run = p_cap.add_run(caption)
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        # Insertar después del heading actual
                        p._p.addnext(p_img._p)
                        p_img._p.addnext(p_cap._p)
                break
    
    print("  + Imágenes de conexiones insertadas con rotulación")

    doc.save(OUT)
    
    # 3. Insertar Capítulo 8 Drenaje antes de "ANEXO A. PLANOS"
    doc = Document(OUT)
    
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip() == "ANEXO A. PLANOS":
            anchor_idx = i
            break
    if anchor_idx is not None:
        aux = Document()
        _h(aux, "8. DRENAJE PLUVIAL", 1)
        _para(aux, "Documento complementario de drenaje pluvial unificado a la "
                   "memoria de cálculo. Unidades: m², mm/h, L/s, m.")
        _h(aux, "8.1 Resumen ejecutivo", 2)
        _bullet(aux, "Normativa: CE.040 Drenaje Pluvial (RNE RM 126-2021).")
        _bullet(aux, "Método racional: Q = 0.95·C·I·A/3600 (L/s); C = 0.95.")
        _bullet(aux, "Intensidad IILA-SENAMHI-UNI: I ≈ 30.33 mm/h (T = 25 a, "
                      "tc = 10 min, zona I23₁₀, Huancavelica).")
        _bullet(aux, "Área captación: A = 80,295 m² (proyección horizontal).")
        _bullet(aux, "Q_total = 0.643 L/s → 0.161 L/s/c/bajante (4 unidades).")
        _bullet(aux, "Canaleta caja 0.25×0.14 m (1 %), OK por Manning.")
        _bullet(aux, "Bajantes Ø50 mm PVC (n = 0.010): 2 por canalón "
                      "(criterio constructivo).")
        if os.path.exists(DRENAJE):
            _h(aux, "8.2 Texto íntegro del informe", 2)
            from sap2000gen.text_readers import docx_to_text
            for num, t in docx_to_text(DRENAJE):
                if t.strip():
                    _para(aux, t, size=Pt(9))
        else:
            _bullet(aux, f"No se encontró {DRENAJE}.")
        
        anchor = doc.paragraphs[anchor_idx]._p
        nodos = [copy.deepcopy(n) for n in aux._element.body]
        anchor.addprevious(nodos[0])
        for k in range(1, len(nodos)):
            nodos[k - 1].addnext(nodos[k])
        print("  + Capítulo 8 (Drenaje) insertado antes de ANEXO A")

    doc.save(OUT)
    # Limpieza final: eliminar índice interno del drenaje
    n_removed = _limpiar_indice_drenaje(doc)
    print(f"  [OK] Índice interno de drenaje limpiado: {n_removed} párrafos eliminados")
    doc.save(OUT)
    print(f"OK: documento final -> {OUT}")
    return OUT


def _limpiar_indice_drenaje(doc):
    """Elimina el índice interno del drenaje que se cuela en el documento final."""
    import re
    to_remove = []
    in_drenaje_index = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt == "ÍNDICE" or txt == "INDICE":
            in_drenaje_index = True
            to_remove.append(i)
        elif in_drenaje_index:
            txt = p.text.strip()
            # Patrones para entradas de TOC: "1.  TEXTO\tN" o "1.1.  TEXTO\tN"
            # Usamos \s+ para espacios, [^\t]+ para texto sin tabs, \t para tab literal
            if re.match(r'^\d+\.\s+[^\t]+\t\d+$', txt) or re.match(r'^\d+\.\d+\.\s+[^\t]+\t\d+$', txt):
                to_remove.append(i)
            else:
                in_drenaje_index = False
    # Eliminar en orden inverso para no afectar índices
    for i in reversed(to_remove):
        p = doc.paragraphs[i]
        p._p.getparent().remove(p._p)
    return len(to_remove)


if __name__ == "__main__":
    build()
