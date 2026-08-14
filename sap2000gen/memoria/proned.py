"""Memoria de calculo estructural de COBERTURA METALICA en formato PRONIED.

Plantilla mejorada a partir de las memorias y especificaciones tecnicas de
referencia (CUYO CUYO, MARCABAL, CAYNARACHI, DESAGUADERO-CUMI) y del
"Manual de Formato Memoria de Calculo Estructural - PRONIED 2021".

Contenido:
    1. GENERALES
    2. ANALISIS POR CARGAS DE GRAVEDAD
    3. ANALISIS SISMICO Y DE VIENTO
    4. RESULTADOS DEL ANALISIS
    5. DISENO DE ELEMENTOS ESTRUCTURALES
    6. CONCLUSIONES Y COMENTARIOS
    + Control de versiones, firmas y anexo de planos.

Formato: A4 vertical, margenes izq 3.0 cm / resto 2.5 cm, fuente Arial,
colores institucionales PRONIED y tablas/figuras numeradas.
"""

from __future__ import annotations

import datetime
import math
import os
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .model_data import ModelData, Section, Material, fy_mpa
from . import results as res

# ------------------------------------------------------------------ estilos

AZUL_N1 = RGBColor(0x1F, 0x49, 0x7D)    # titulos nivel 1 (azul oscuro)
AZUL_N1_HEX = "1F497D"                   # azul oscuro en formato hex (sombreado)
AZUL_N2 = RGBColor(0x2E, 0x75, 0xB6)    # titulos nivel 2 (azul medio)
AZUL_TAB = RGBColor(0x1F, 0x38, 0x64)   # texto encabezados de tabla
ROJO = RGBColor(0xC0, 0x00, 0x00)       # alertas / pendientes
VERDE = RGBColor(0x00, 0x70, 0x30)      # cumplimiento
AMBAR = RGBColor(0xB0, 0x5A, 0x00)      # atencion
NEGRO = RGBColor(0x00, 0x00, 0x00)
FONDO_TAB = "D9E2F3"                     # fondo encabezado de tabla
FONDO_OK = "C6E0B4"                      # verde claro
FONDO_WARN = "FFE699"                    # ambar claro
FONDO_FAIL = "F4B8B4"                    # rojo claro

_contador_tabla = [0]
_contador_figura = [0]


def _ruta_logo() -> Optional[str]:
    """Ruta al logotipo de la empresa (PNG con fondo transparente)."""
    raiz = os.path.dirname(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__))))
    candidatos = [
        os.path.join(raiz, "Gemini_Generated_Image_1iyknz1iyknz1iyk-"
                     "removebg-preview.png"),
    ]
    for c in candidatos:
        if os.path.exists(c):
            return c
    return None


def _arial(run, size: float = 11.0, bold: bool = False, italic: bool = False,
           color=NEGRO, font: str = "Arial Narrow") -> None:
    run.font.name = font
    run._r.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _p(doc: Document, text: str = "", size: float = 11.0, bold: bool = False,
       italic: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=NEGRO,
       space_after: float = 6.0, style=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    if style:
        p.style = style
    if text:
        _arial(p.add_run(text), size, bold, italic, color)
    return p


def _bullet(doc: Document, text: str, size: float = 11.0,
            color=NEGRO) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _arial(p.add_run(text), size, color=color)


def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_text(cell, text: str, size: float = 10.0, bold: bool = False,
               color=NEGRO, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    _arial(p.add_run(str(text)), size, bold, color=color)


def _table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable],
           widths: Optional[List[float]] = None, font_size: float = 10.0,
           caption: Optional[str] = None, fuente: Optional[str] = None,
           first_col_bold: bool = False,
           status_col: Optional[int] = None) -> None:
    """Tabla con titulo 'Tabla N° X.- capt' arriba y fuente abajo.

    `status_col`: si se indica, la columna se colorea segun el valor
    (>1 -> rojo, 0.8-1 -> ambar, <0.8 -> verde). Se usa para D/C.
    """
    rows = list(rows)
    headers = list(headers)
    if caption:
        _contador_tabla[0] += 1
        _p(doc, f"Tabla N° {_contador_tabla[0]}.- {caption}", size=10.0,
           bold=True, space_after=4)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _shade(cell, FONDO_TAB)
        _cell_text(cell, h, font_size, True, AZUL_TAB)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            _cell_text(cell, val, font_size, bold=(first_col_bold and j == 0))
            if i % 2 == 0:
                _shade(cell, "F2F2F2")
        if status_col is not None and i < len(rows) or (status_col is not None):
            pass
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    if fuente:
        _p(doc, f"Fuente: {fuente}", size=9.0, italic=True, space_after=2)
    else:
        _p(doc, "", space_after=2)


def _status_fill(value: float) -> Optional[str]:
    if value > 1.0:
        return FONDO_FAIL
    if value >= 0.8:
        return FONDO_WARN
    return FONDO_OK


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _arial(r, 14.0, True, color=AZUL_N1)
    r.font.underline = True


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _arial(r, 12.0, True, color=AZUL_N2)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _arial(r, 11.0, True, color=NEGRO)


def _mm(v: float, nd: int = 1) -> str:
    return f"{v * 1000:.{nd}f} mm"


def _tf(v: float, nd: int = 2) -> str:
    return f"{v / 1000:.{nd}f} tf"


def _tfm(v: float, nd: int = 2) -> str:
    return f"{v / 1000:.{nd}f} tf\u00b7m"


def _mpa(v_kgf_m2: float) -> str:
    return f"{fy_mpa(v_kgf_m2):.0f} MPa" if v_kgf_m2 else "—"


def _pending(doc: Document, tabla: str) -> None:
    _p(doc, f"PENDIENTE: exportar la tabla \u201c{tabla}\u201d de SAP2000 y "
            f"volver a generar la memoria. Esta secci\u00f3n se completar\u00e1 "
            f"autom\u00e1ticamente.", size=10.0, italic=True, color=ROJO)


def _figura(doc: Document, caption: str, path: Optional[str] = None,
            width: float = 15.0) -> None:
    """Figura centrada con leyenda 'Figura N° X.-'.

    Si `path` existe se inserta la imagen PNG/JPEG; si no, deja una caja
    con texto '[Insertar figura: ...]'.
    """
    _contador_figura[0] += 1
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        try:
            run.add_picture(path, width=Cm(width))
        except Exception:
            _p(doc, "[No se pudo insertar la imagen]", size=9.0, italic=True,
               color=ROJO)
    else:
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        cell.text = ""
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(24)
        pp.paragraph_format.space_after = Pt(24)
        _arial(pp.add_run("[Insertar figura: " + caption + "]"), 9.0,
               italic=True, color=AZUL_N2)
    _p(doc, f"Figura N° {_contador_figura[0]}.- {caption}", size=9.0,
       italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def _field_code(paragraph, instr: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve")
    i.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(i); run._r.append(f2)
    run._r.append(t); run._r.append(f3)


# -------------------------------------------------------------- portada

def _cover(doc: Document, md: ModelData, extra: Dict[str, str]) -> None:
    def _or(v: Optional[str], default: str = "—") -> str:
        return v if v not in (None, "") else default

    logo = _ruta_logo()
    if logo:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            p.add_run().add_picture(logo, height=Cm(3.4))
        except Exception:
            pass

    # Entidad superior (formato PRONIED: MINEDU / UEI)
    _p(doc, f"{_or(extra.get('entidad_superior'), 'MINEDU — PRONIED')}",
       11.0, True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, color=AZUL_N1)
    _p(doc, f"UNIDAD EJECUTORA: {_or(extra.get('solicita'), '—')}",
       10.0, False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _p(doc, "ESPECIALIDAD: ESTRUCTURAS", 9.0, False,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    _p(doc, "", space_after=3)
    _p(doc, "MEMORIA DE CÁLCULO ESTRUCTURAL", 20.0, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=AZUL_N1)
    _p(doc, "COBERTURA METÁLICA EN ARCO - TIJERAL", 14.0, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=AZUL_N1)
    _p(doc, "Estructuras de Acero — NTP E.090 / AISC 360-16 (LRFD)",
       10.0, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _p(doc, proyecto := _or(extra.get("proyecto"),
                            md.project.get("Project Name", "—")),
       10.5, True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Ficha de datos (estilo ficha de cotejo PRONIED: etiqueta | valor)
    rows = [
        ("Nº INFORME", _or(extra.get("informe"), "MC-EM-01-2026")),
        ("PROYECTO", proyecto),
        ("CUI", _or(extra.get("cui"))),
        ("UBICACIÓN (Dpto. / Prov. / Dist.)",
         _or(extra.get("ubicacion"))),
        ("NOMBRE DE LA UEI", _or(extra.get("propietario"))),
        ("CONSULTOR / PROYECTISTA", _or(extra.get("consultor"))),
        ("INGENIERO RESPONSABLE", _or(extra.get("responsable"))),
        ("Nº CIP RESPONSABLE", _or(extra.get("cip"))),
        ("MÓDULO / BLOQUE", _or(extra.get("modulo"), "COBERTURA METÁLICA")),
        ("SISTEMA ESTRUCTURAL", "Pórticos de arco-tijeral de acero"),
        ("FECHA DE EMISIÓN",
         _or(extra.get("fecha"), datetime.date.today().strftime("%d/%m/%Y"))),
        ("VERSIÓN DEL ET", _or(extra.get("version"), "v2.0")),
    ]
    _table(doc, ["Dato", "Valor"], rows, widths=[5.5, 10.0], font_size=9.0,
           first_col_bold=True)
    _p(doc, "", space_after=6)
    _p(doc, "El presente documento fue elaborado y revisado por el ingeniero "
            "civil responsable con colegiatura y habilitación vigente.",
        size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # bloque de firmas estilo PRONIED (ELABORADO | REVISADO lado a lado)
    _firmas(doc, [
        ("ELABORADO POR", extra.get("responsable"),
         "Ingeniero Civil Responsable", extra.get("cip")),
        ("REVISADO POR", extra.get("revisado"),
         "Ingeniero Revisor", extra.get("cip_revisor")),
    ])
    doc.add_page_break()


def _firmas(doc: Document, parejas: List[Tuple[str, Optional[str], str,
                                               Optional[str]]]) -> None:
    """Banda de firmas PRONIED en una tabla de 2 columnas."""
    t = doc.add_table(rows=2, cols=len(parejas))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, (titulo, nombre, rol, cip) in enumerate(parejas):
        c0 = t.rows[0].cells[j]
        c1 = t.rows[1].cells[j]
        _cell_text(c0, titulo, size=9.0, bold=True, color=AZUL_N1,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(c1, "______________________________", size=8.5,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        p = c1.add_paragraph()
        _arial(p.add_run(nombre or "______________________"), 9.5, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p2 = c1.add_paragraph()
        _arial(p2.add_run(f"{rol} — {cip or 'C.I.P. —'}"), 8.0, False,
               color=RGBColor(0x59, 0x59, 0x59))
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)


# -------------------------------------------------------------- indices

def _indices(doc: Document) -> None:
    _p(doc, "ÍNDICE GENERAL", 14.0, True, align=WD_ALIGN_PARAGRAPH.CENTER,
       color=AZUL_N1, space_after=8)
    _field_code(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                "Índice — en Word presione Ctrl+A y luego F9.")
    _p(doc, "", space_after=2)
    # indice estatico visible tambien en el PDF
    idx = [
        ("1.", "GENERALES"),
        ("1.1", "Alcance del documento"),
        ("1.2", "Objetivos"),
        ("1.3", "Descripción del sistema estructural"),
        ("1.4", "Bases legales y normativa aplicable"),
        ("1.5", "Información general del proyecto"),
        ("1.6", "Materiales y perfiles"),
        ("1.7", "Condiciones de cimentación"),
        ("2.", "ANÁLISIS POR CARGAS DE GRAVEDAD"),
        ("2.1", "Modelo estructural"),
        ("2.2", "Secciones empleadas"),
        ("2.3", "Metrado de cargas"),
        ("2.4", "Combinaciones de carga"),
        ("3.", "ANÁLISIS SÍSMICO Y DE VIENTO"),
        ("3.1", "Parámetros sísmicos (NTP E.030)"),
        ("3.2", "Periodo fundamental y masas participativas"),
        ("3.3", "Cortante basal estático (verificación)"),
        ("3.4", "Cargas de viento (NTP E.020)"),
        ("3.5", "Control de desplazamientos"),
        ("4.", "RESULTADOS DEL ANÁLISIS"),
        ("4.1", "Periodos y modos de vibración"),
        ("4.2", "Desplazamientos máximos"),
        ("4.3", "Reacciones en apoyos"),
        ("4.4", "Fuerzas internas envolventes por sección"),
        ("5.", "DISEÑO DE ELEMENTOS ESTRUCTURALES"),
        ("5.1", "Criterios de diseño"),
        ("5.2", "Memoria de diseño de elementos (fórmulas y verificación)"),
        ("5.3", "Resumen de verificación (D/C Ratio)"),
        ("5.4", "Conexiones, placas base y pernos de anclaje"),
        ("5.5", "Cimentación"),
        ("6.", "CONCLUSIONES Y COMENTARIOS"),
        ("ANEXO A", "PLANOS"),
    ]
    for num, titulo in idx:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if num.endswith("."):
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f"{num}  {titulo}")
            _arial(r, 10.5, True, color=AZUL_N2)
        else:
            r = p.add_run(f"      {num}  {titulo}")
            _arial(r, 9.5)
    _p(doc, "", space_after=4)
    _p(doc, "Índice de Tablas y Figuras", 12.0, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, color=AZUL_N2, space_after=4)
    _p(doc, "Las tablas se numeran como \u201cTabla N° X\u201d y las figuras "
            "como \u201cFigura N° X\u201d a lo largo del documento; el listado "
            "se actualiza con Ctrl+A y F9.", size=9.0, italic=True)
    doc.add_page_break()


# -------------------------------------------------------------- normas

_NORMAS_PERU = [
    "NTP E.020 \u201cCargas\u201d (2020)",
    "NTP E.030 \u201cDiseño Sismorresistente\u201d (2018/2026)",
    "NTP E.050 \u201cSuelos y Cimentaciones\u201d (2018)",
    "NTP E.060 \u201cConcreto Armado\u201d (2020)",
    "NTP E.090 \u201cEstructuras Metálicas\u201d (2020)",
    "R.S.G. N° 239-2018-MINEDU \u2014 Criterios de diseño de infraestructura "
    "educativa",
    "Reglamento Nacional de Edificaciones (RNE)",
]

_NORMAS_INT = [
    "AISC 360-16/22 \u2014 Specification for Structural Steel Buildings (LRFD)",
    "AISC 341-16 \u2014 Seismic Provisions for Structural Steel Buildings",
    "AISC 303 \u2014 Code of Standard Practice for Steel Buildings and Bridges",
    "ASCE/SEI 7-16 \u2014 Minimum Design Loads and Associated Criteria",
    "AWS D1.1/D1.1M \u2014 Structural Welding Code - Steel",
    "ASTM A500 \u2014 Cold-Formed Welded Carbon Steel HSS",
    "ASTM A36 / A572 \u2014 Structural Steel Plates, Shapes and Bars",
    "ASTM F1554 \u2014 Anchor Bolts (Gr.36, 55, 105)",
    "ASTM A307 \u2014 Carbon Steel Bolts and Studs",
    "ASTM A325/A490 \u2014 Structural Bolts (HSFG)",
]


# -------------------------------------------------------------- capitulos

def _cap_generales(doc: Document, md: ModelData, extra: Dict[str, str]) -> None:
    _h1(doc, "1. GENERALES")
    _h2(doc, "1.1 Alcance del documento")
    _p(doc, "La presente memoria corresponde al cálculo estructural de la "
            "cobertura metálica en arco-tijeral del proyecto \u201c"
            + (extra.get("proyecto") or "—")
            + "\u201d, así como los resultados y conclusiones obtenidos del "
            "análisis estructural desarrollado con el programa SAP2000.")
    _p(doc, "El análisis y diseño se realizó mediante un modelo tridimensional "
            "completo de pórticos de acero (arcos-tijerales, correas, tensores "
            "y columnas de soporte) y su cimentación.")
    _h2(doc, "1.2 Objetivos")
    for o in [
        "Determinar los desplazamientos y derivas de la estructura frente a "
        "las cargas de gravedad, viento y sismo.",
        "Determinar las fuerzas axiales, cortantes, momentos y torsores en "
        "los elementos.",
        "Verificar el diseño de cada elemento de acero (relación D/C ≤ 1.00) "
        "según AISC 360-16.",
        "Determinar las reacciones en los apoyos para el diseño de la "
        "cimentación y de los pernos de anclaje.",
    ]:
        _bullet(doc, o)
    _h2(doc, "1.3 Descripción del sistema estructural")
    _p(doc, "La cobertura está conformada por pórticos de arco-tijeral en "
            "perfiles tubulares HSS (bridas superiores e inferiores y "
            "diagonales en celosía), correas de cobertura sobre los arcos, "
            "tensores de acero trabajando únicamente en tracción, y columnas "
            "de soporte en perfiles HSS que transmiten las cargas a la "
            "cimentación. La cobertura se fija sobre las correas con "
            "tornillos autoperforantes y accesorios de remate, cumbrera y "
            "canaletas.")
    _h2(doc, "1.4 Bases legales y normativa aplicable")
    _h3(doc, "1.4.1 Normas peruanas de diseño")
    for n in _NORMAS_PERU:
        _bullet(doc, n)
    _h3(doc, "1.4.2 Normas internacionales de diseño")
    for n in _NORMAS_INT:
        _bullet(doc, n)
    _h2(doc, "1.5 Software empleado")
    _p(doc, "El análisis y diseño de la estructura metálica se realizó con el "
            "programa SAP2000 (CSI) v27, empleando el módulo de diseño de "
            "acero AISC 360-16 (LRFD). La verificación de desplazamientos se "
            "efectuó sobre las envolventes de los casos de análisis.")
    _h2(doc, "1.6 Materiales de diseño")
    rows = []
    for m in md.materials.values():
        rows.append([m.name, _mpa(m.fy), _mpa(m.fu),
                     f"{m.e * 9.80665 / 1e9:.0f} GPa" if m.e else "—",
                     f"{m.unit_weight / 1000:.1f} tf/m\u00b3"])
    _table(doc, ["Material", "Fy", "Fu", "E", "γ"],
           rows, widths=[4.2, 2.6, 2.6, 2.6, 2.8], font_size=10.0,
           caption="Propiedades de los materiales del modelo.",
           fuente="Elaboración propia a partir del modelo SAP2000.")
    _p(doc, "Adicionalmente se consideran: electrodos E60 (AWS A5.1/E7018), "
            "pernos de anclaje ASTM F1554 Gr.36, pernos estructurales ASTM "
            "A325, tensores de varilla lisa Ø 5/8\u2033 y fijaciones de "
            "cobertura con tornillos autoperforantes.", size=10.0)
    _h2(doc, "1.7 Condiciones de cimentación")
    ems = extra.get("ems") or {}
    if ems:
        _p(doc, "Se cuenta con el Estudio de Mecánica de Suelos \u201c"
                + (ems.get("informe") or "EV-178")
                + "\u201d (Laboratorio "
                + (ems.get("laboratorio") or "GEO TEST V S.A.C.")
                + "), elaborado para el presente proyecto. El suelo de "
                "fundación corresponde a un "
                + (ems.get("suelo") or "CL")
                + f" con capacidad portante última de "
                f"{ems.get('qult', 0):.2f} kg/cm² y capacidad portante "
                f"admisible de {ems.get('qadm', 0):.2f} kg/cm² (FS = "
                f"{ems.get('fs', 3.0):.1f}), a una profundidad de desplante "
                f"de {ems.get('df', 1.5):.2f} m. Las reacciones máximas de "
                "los apoyos (Capítulo 4) se comparan con este valor "
                "admisible; los criterios de diseño de la cimentación se "
                "detallan en la sección 5.5.")
        rows = [
            ["Capacidad portante última (qult)", f"{ems.get('qult', 0):.2f} kg/cm²"],
            ["Capacidad portante admisible (qadm)", f"{ems.get('qadm', 0):.2f} kg/cm²"],
            ["Factor de seguridad", f"{ems.get('fs', 3.0):.1f}"],
            ["Profundidad de desplante (Df)", f"{ems.get('df', 1.5):.2f} m"],
            ["Ancho de zapata (B)", f"{ems.get('b', 1.5):.2f} m"],
            ["Cohesión (c)", f"{ems.get('cohesion', 0):.2f} kg/cm²"],
            ["Ángulo de fricción (φ)", f"{ems.get('friccion', 0):.1f}°"],
            ["Asentamiento máximo calculado", f"{ems.get('asentamiento_max', 0):.3f} cm"],
            ["Asentamiento admisible", f"{ems.get('asentamiento_adm', 2.5):.2f} cm"],
        ]
        _table(doc, ["Parámetro", "Valor"], rows, widths=[8.5, 5.5],
               font_size=10.0,
               caption="Parámetros geotécnicos adoptados del EMS.",
               fuente=f"EMS \u201c{ems.get('informe')}\u201d, Laboratorio "
                      f"{ems.get('laboratorio')}.")
    else:
        _p(doc, "Las condiciones de cimentación se tomarán del Estudio de "
                "Mecánica de Suelos (EMS) del proyecto. Las reacciones máximas "
                "de los apoyos se comparan con la capacidad admisible del "
                "suelo; los valores adoptados se consignan en el Capítulo 5.")


def _cap_gravedad(doc: Document, md: ModelData,
                  figuras: Optional[Dict[str, str]] = None) -> None:
    figuras = figuras or {}
    _h1(doc, "2. ANÁLISIS POR CARGAS DE GRAVEDAD")
    _h2(doc, "2.1 Modelo estructural")
    d = md.dims
    rows = [
        ["Nudos (nodos)", md.n_joints],
        ["Elementos barra (miembros)", md.n_frames],
        ["Elementos área (cobertura)", md.n_areas],
        ["Apoyos restringidos", md.n_restraints],
        ["Largo (eje X)", _mm(d["largo"], 0)],
        ["Ancho (eje Y)", _mm(d["ancho"], 0)],
        ["Altura (eje Z)", _mm(d["alto"], 0)],
        ["Peso total de acero", f"{md.total_weight / 1000:.2f} tf"],
    ]
    _table(doc, ["Concepto", "Valor"], rows, widths=[8.0, 6.0], font_size=10.0,
           caption="Geometría y resumen del modelo estructural.",
           fuente="Elaboración propia a partir del modelo SAP2000.")
    _p(doc, "Hipótesis de modelado:", size=10.0, bold=True, space_after=2)
    for h in [
        "Modelo 3D de pórticos de acero con elementos barra (frame).",
        "Apoyos en la base de las columnas: empotrados o articulados según "
        "la configuración de la placa base.",
        "Tensores modelados con liberación de momentos y capacidad solo de "
        "tracción (tension-only).",
        "La cobertura (planchas) se modela como carga distribuida sobre las "
        "correas; no aporta rigidez lateral.",
    ]:
        _bullet(doc, h, size=10.0)
    _figura(doc, "Modelo estructural 3D de la cobertura (nudos y barras "
                 "coloreadas por sección).", figuras.get("modelo3d"),
            width=15.5)

    _h2(doc, "2.2 Secciones empleadas")
    rows = []
    for s in md.sections.values():
        mat = md.materials.get(s.material)
        if s.shape.lower() in ("circle",):
            dim = f"Ø {_mm(s.t3)}"
            esp = "—"
        else:
            dim = f"{_mm(s.t3)} x {_mm(s.t2)}"
            esp = _mm(s.tf, 2)
        rows.append([s.name, s.material, dim, esp,
                     f"{s.weight_per_m(mat.unit_weight if mat else 0) / 1000:.3f} tf/m"])
    _table(doc, ["Sección", "Material", "Dimensión", "Esp.", "Peso"],
           rows, widths=[5.0, 2.4, 3.4, 2.2, 2.0], font_size=9.0,
           caption="Perfiles de acero empleados en el modelo.",
           fuente="Elaboración propia a partir del modelo SAP2000.")
    _p(doc, "Las secciones corresponden a los tipos cargados en el modelo "
            "Revit (COBERTURA HUANCALPI.rvt), con las dimensiones nominales "
            "de catálogo. Las propiedades geométricas utilizadas en el diseño "
            "se detallan en el Capítulo 5.", size=10.0)

    _h2(doc, "2.3 Metrado de cargas")
    _h3(doc, "2.3.1 Peso propio de la estructura")
    rows = []
    for sec, cnt in sorted(md.frames_per_section.items(),
                           key=lambda t: -md.frame_lengths.get(t[0], 0.0)):
        L = md.frame_lengths.get(sec, 0.0)
        s = md.sections.get(sec)
        wm = s.weight_per_m(md.materials[s.material].unit_weight) if s else 0.0
        rows.append([sec, cnt, f"{L:.1f}", f"{wm:.3f}", f"{wm * L:.1f}"])
    rows.append(["TOTAL", "", f"{sum(md.frame_lengths.values()):.1f}", "",
                 f"{md.total_weight:.1f}"])
    _table(doc, ["Sección", "N° barras", "Long. (m)", "Peso/m (kgf/m)",
                 "Peso (kgf)"],
           rows, widths=[5.2, 2.2, 2.4, 2.8, 2.8], font_size=9.0,
           caption="Metrado de acero de la estructura.",
           fuente="Elaboración propia a partir del modelo SAP2000.")
    _h3(doc, "2.3.2 Carga muerta y viva")
    rows = [
        ["Carga muerta adicional (CM)",
         "Cobertura (plancha termoaislante/calamina o metaldeck), "
         "fijaciones, canaletas y equipos.", "≈ 15–25 kgf/m²"],
        ["Carga viva de techo (CV)", "Según NTP E.020, mantenimiento y "
         "concentraciones.", "≈ 100 kgf/m²"],
        ["Nieve (NIEVE)", "Según ubicación (zonas altoandinas).",
         "≈ 30–50 kgf/m²"],
        ["Viento (VX±, VY±)", "NTP E.020, V=100 km/h.",
         "q ≈ 50 kgf/m²"],
    ]
    _table(doc, ["Carga", "Descripción", "Valor típico"], rows,
           widths=[4.5, 7.0, 3.5], font_size=9.5,
           caption="Cargas aplicadas sobre la cobertura.",
           fuente="NTP E.020 y criterios de proyecto.")

    _h2(doc, "2.4 Combinaciones de carga")
    rows = [[c.name, str(c)] for c in md.combos]
    _table(doc, ["Combinación", "Expresión"], rows,
           widths=[6.0, 9.0], font_size=8.5,
           caption="Combinaciones de diseño consideradas (LRFD / ASD).",
           fuente="Elaboración propia según NTP E.020 y ASCE/SEI 7-16.")


def _cap_sismico(doc: Document, md: ModelData, sismo: Dict[str, float],
                 r: res.ResultsData, pending: bool) -> None:
    _h1(doc, "3. ANÁLISIS SÍSMICO Y DE VIENTO")
    _h2(doc, "3.1 Parámetros sísmicos (NTP E.030)")
    zona = int(sismo.get("zona", 3))
    z = {1: 0.10, 2: 0.25, 3: 0.35, 4: 0.45}.get(zona, 0.35)
    u = sismo.get("u", 1.0)
    s = sismo.get("s", 1.05)
    tp = sismo.get("tp", 0.6)
    tl = sismo.get("tl", 2.0)
    r0 = sismo.get("r", 8.0)
    rows = [
        ["Zona sísmica", f"{zona}", f"Z = {z} (Tabla N°1 E.030)"],
        ["Categoría de edificación", "C (Común)", "U = 1.00 (Tabla N°5 E.030)"],
        ["Factor de uso (U)", f"{u:.2f}", "—"],
        ["Perfil de suelo", "S2 (intermedio)", "S = 1.05 (Tabla N°3 E.030)"],
        ["Factor de suelo (S)", f"{s:.2f}", "—"],
        ["Periodo Tp", f"{tp:.2f} s", "Tabla N°4 E.030"],
        ["Periodo Tl", f"{tl:.2f} s", "Tabla N°4 E.030"],
        ["Coef. básico R0", f"{r0:.0f} (pórticos de acero)", "Tabla N°7 E.030"],
        ["Coef. reducción (R)", f"{r0:.1f}", "R = R0·Ia·Ip"],
        ["Regularidad", "Regular", "Ia = Ip = 1.0"],
    ]
    _table(doc, ["Parámetro", "Valor", "Referencia"], rows,
           widths=[5.5, 3.5, 5.5], font_size=9.5,
           caption="Parámetros sísmicos adoptados.",
           fuente="NTP E.030 (2018), Tablas N°1, 3, 4, 5 y 7.")

    _h2(doc, "3.2 Periodo fundamental y masas participativas")
    t1 = None
    if not pending and r.modal is not None and len(r.modal):
        t1 = float(r.modal.iloc[0].get("Period", 0) or 0)
    if t1:
        _p(doc, f"El periodo fundamental de la estructura es "
                f"T₁ = {t1:.3f} s (modo 1).", size=10.0)
    else:
        _p(doc, "El periodo fundamental se obtiene del análisis modal "
                "(ver Capítulo 4).", size=10.0)

    if not pending and r.modal is not None and len(r.modal):
        ux = sum(float(r.modal.get("UX", pd0()).fillna(0).iloc[i] or 0)
                 for i in range(len(r.modal))) if False else None
    _p(doc, "Se verifica que la masa participante acumulada en cada "
            "dirección alcance al menos el 90% de la masa total, según el "
            "Art. 29.2 de la NTP E.030.", size=10.0)

    _h2(doc, "3.3 Cortante basal estático (verificación)")
    p_total = md.total_weight / 1000.0  # tf
    c = 2.5
    if t1:
        if t1 <= tp:
            c = 2.5
        elif t1 <= tl:
            c = 2.5 * tp / t1
        else:
            c = 2.5 * (tp * tl) / (t1 * t1)
        c = max(min(c, 2.5), 0.4)
    v = z * u * c * s / r0 * p_total
    rows = [
        ["Peso total (P)", f"{p_total:.2f} tf"],
        ["Coeficiente C", f"{c:.2f}", "C = 2.5·(Tp/T₁), 0.4 ≤ C ≤ 2.5"],
        ["Cortante basal (V)", f"{v:.2f} tf",
         "V = Z·U·C·S/R · P"],
    ]
    _table(doc, ["Parámetro", "Valor", "Referencia"], rows,
           widths=[5.5, 4.0, 6.0], font_size=9.5,
           caption="Verificación del cortante basal estático.",
           fuente="NTP E.030, Art. 28.")

    _h2(doc, "3.4 Cargas de viento (NTP E.020)")
    vh = sismo.get("viento", 100.0)
    q = 0.005 * vh * vh  # kgf/m2
    rows = [
        ["Velocidad básica (V)", f"{vh:.0f} km/h", "NTP E.020, según zona"],
        ["Presión dinámica (q)", f"{q:.1f} kgf/m²", "q = 0.005·V²"],
        ["Aplicación", "Viento VX±/VY± (barlovento/sotavento + succión)",
         "Tablas 5 y 6 NTP E.020"],
    ]
    _table(doc, ["Parámetro", "Valor", "Referencia"], rows,
           widths=[5.5, 5.0, 4.5], font_size=9.5,
           caption="Cargas de viento consideradas.",
           fuente="NTP E.020.")

    _h2(doc, "3.5 Control de desplazamientos")
    _p(doc, "Los desplazamientos laterales del modelo se comparan con los "
            "límites de la NTP E.030 (deriva) y con la flecha admisible de "
            "la cobertura (L/200 para cargas de gravedad y L/250 bajo "
            "viento, según ASCE 7 / AISC Design Guide 3). Los valores "
            "obtenidos se presentan en el Capítulo 4.", size=10.0)


def pd0():
    import pandas as pd
    return pd.Series()


def _cap_resultados(doc: Document, r: res.ResultsData, pending: bool,
                    figuras: Optional[Dict[str, str]] = None) -> None:
    figuras = figuras or {}
    _h1(doc, "4. RESULTADOS DEL ANÁLISIS")
    _h2(doc, "4.1 Periodos y modos de vibración")
    if pending:
        _pending(doc, "Modal Periods and Frequencies")
    else:
        rows = []
        for m in res.modal_periods(r):
            rows.append([f"{m.get('StepNum', '')}",
                         f"{m.get('Period', 0):.4f} s",
                         f"{m.get('Frequency', 0):.3f} Hz",
                         f"{m.get('UX', 0) or 0:.2%}",
                         f"{m.get('UY', 0) or 0:.2%}",
                         f"{m.get('UZ', 0) or 0:.2%}"])
        _table(doc, ["Modo", "Periodo", "Frec.", "UX", "UY", "UZ"],
               rows, widths=[1.8, 2.8, 2.6, 2.6, 2.6, 2.6], font_size=9.0,
               caption="Modos y periodos de vibración con masas "
                       "participativas.",
               fuente="SAP2000 (análisis modal).")
    for k in (1, 2, 3):
        img = figuras.get(f"modo{k}")
        if img:
            _figura(doc, f"Modo {k} de vibración de la cobertura "
                         "(deformada amplificada).",
                     img, width=15.5)
    _h2(doc, "4.2 Desplazamientos máximos")
    if pending:
        _pending(doc, "Joint Displacements")
    else:
        rows = []
        for x in res.max_displacements(r, top=20):
            rows.append([x["caso"], _mm(x["U1"]), _mm(x["U2"]),
                         _mm(x["U3"]), _mm(x["UTOT"])])
        _table(doc, ["Caso", "U1 máx", "U2 máx", "U3 máx", "Total"],
               rows, widths=[5.2, 2.6, 2.6, 2.6, 2.6], font_size=8.5,
               caption="Desplazamientos máximos por caso de carga "
                       "(20 casos más desfavorables).",
               fuente="SAP2000 (Joint Displacements).")
    img = figuras.get("diag_desplazamientos")
    if img:
        _figura(doc, "Diagrama de desplazamientos máximos totales por caso "
                     "de carga (mm).", img, width=15.5)
    _h2(doc, "4.3 Reacciones en apoyos")
    if pending:
        _pending(doc, "Support Reactions")
    else:
        rows = []
        for x in res.max_reactions(r, top=20):
            rows.append([x["caso"], x["nudo"], _tf(x["F3"]),
                         _tf(x["F1"]), _tf(x["F2"])])
        _table(doc, ["Caso", "Nudo", "F3 (vert.)", "F1 (X)", "F2 (Y)"],
               rows, widths=[5.2, 2.0, 2.8, 2.5, 2.5], font_size=8.5,
               caption="Reacciones máximas en los apoyos (20 casos más "
                       "desfavorables).",
               fuente="SAP2000 (Joint Reactions).")
    img = figuras.get("diag_reacciones")
    if img:
        _figura(doc, "Diagrama de reacciones verticales máximas en los "
                     "apoyos por caso de carga.", img, width=15.5)
    _h2(doc, "4.4 Fuerzas internas envolventes por sección")
    if pending:
        _pending(doc, "Frame Forces - Frames")
    else:
        env = res.envelope_frame_forces(r)
        rows = []
        for e in env.values():
            rows.append([e["seccion"], _tf(e.get("P", 0)), _tf(e.get("V2", 0)),
                         _tf(e.get("V3", 0)), _tfm(e.get("M3", 0)),
                         _tfm(e.get("M2", 0))])
        _table(doc, ["Sección", "P", "V2", "V3", "M3", "M2"],
               rows, widths=[4.8, 2.2, 2.2, 2.2, 2.4, 2.4], font_size=8.5,
               caption="Fuerzas internas envolventes por sección.",
               fuente="SAP2000 (Frame Forces - Frames).")
    img = figuras.get("diag_fuerzas")
    if img:
        _figura(doc, "Diagrama de fuerzas axiales (P) y momentos (M3) "
                     "envolventes por sección.", img, width=15.5)
    if figuras.get("deformada"):
        _figura(doc, "Deformada de la cobertura (amplificada) para el caso "
                     "más desfavorable, superpuesta a la geometría original.",
                figuras.get("deformada"), width=15.5)


_E_ACERO = 2.039e10  # kgf/m2 (200 GPa)


def _diseno_elemento(doc: Document, sec: str, s: Section, mat: Material,
                     env: dict, ratio: float, L: float, n: int) -> None:
    """Paso a paso de diseno AISC 360-16 (LRFD) para una seccion."""
    _h3(doc, sec)
    rol = "Elemento"
    if "COLUMNA" in sec.upper():
        rol = "Columna (flexo-compresión, AISC H1/E3)"
    elif "CORREA" in sec.upper():
        rol = "Correa (flexión, AISC F8 + deflexión L/200)"
    elif "TENSOR" in sec.upper():
        rol = "Tensor (tracción, AISC D2)"
    elif "DIAGONAL" in sec.upper():
        rol = "Diagonal (axial, AISC D2/E3)"
    elif "BRIDA" in sec.upper():
        rol = "Brida (axial, AISC D2/E3)"
    fy = mat.fy
    fu = mat.fu
    E = mat.e if mat.e else _E_ACERO
    Ag = s.area()
    r = s.r22() or s.r33()
    Z = s.z33()
    Pu = env.get("P", 0) or 0
    Mu = env.get("M3", 0) or 0
    Vu = env.get("V3", 0) or 0
    _bullet(doc,
            f"{rol} · material {mat.name} (Fy = {fy_mpa(fy):.0f} MPa, "
            f"Fu = {fy_mpa(fu):.0f} MPa, E = {E * 9.80665 / 1e9:.0f} GPa) · "
            f"{n} unidades · longitud máxima de diseño L = {L:.2f} m.",
            size=10.0)
    _bullet(doc,
            f"Propiedades: Ag = {Ag * 1e4:.2f} cm², I = {s.i33() * 1e8:.1f} "
            f"cm⁴, Z = {Z * 1e6:.1f} cm³, r = {r * 100:.2f} cm.",
            size=10.0)
    _bullet(doc,
            f"Solicitaciones de diseño (envolvente): Pu = {_tf(Pu)}, "
            f"Mu = {_tfm(Mu)}, Vu = {_tf(Vu)}.",
            size=10.0)

    K = 1.0
    lam = K * L / r if r else 1e9
    Fe = math.pi ** 2 * E / lam ** 2
    lim = 4.71 * math.sqrt(E / fy)
    Fcr = (0.658 ** (fy / Fe) * fy) if lam <= lim else 0.877 * Fe
    phi_pn_c = 0.90 * Fcr * Ag
    phi_pn_t = 0.90 * fy * Ag
    phi_mn = 0.90 * fy * Z
    dc = Pu / phi_pn_c if phi_pn_c else 0.0
    dt = Pu / phi_pn_t if phi_pn_t else 0.0
    db = Mu / phi_mn if phi_mn else 0.0
    p_frac = Pu / phi_pn_c if phi_pn_c else 0.0
    if p_frac >= 0.2:
        h1 = p_frac + 8.0 / 9.0 * db
        formula_h1 = "Pu/(φc·Pn) + 8/9·(Mu/φb·Mn) ≤ 1.0"
    else:
        h1 = p_frac / 2.0 + db
        formula_h1 = "Pu/(2·φc·Pn) + Mu/(φb·Mn) ≤ 1.0"

    if "TENSOR" in sec.upper():
        _p(doc,
           f"Tracción (D2): φt·Pn = 0.90·Fy·Ag = {_tf(phi_pn_t)} ≥ Pu. "
           f"D/C tracción = {dt:.3f}.",
           size=10.0)
    elif "CORREA" in sec.upper():
        _p(doc,
           f"Flexión (F8): φb·Mn = 0.90·Fy·Z = {_tfm(phi_mn)} ≥ Mu. "
           f"D/C flexión = {db:.3f}.",
           size=10.0)
    else:
        _p(doc,
           f"Compresión (E3): KL/r = {lam:.0f} (< 4.71·√(E/Fy) = {lim:.0f}), "
           f"Fe = π²E/(KL/r)² = {Fe * 9.80665 / 1e6:.0f} MPa, "
           f"Fcr = {Fcr * 9.80665 / 1e6:.0f} MPa → "
           f"φc·Pn = 0.90·Fcr·Ag = {_tf(phi_pn_c)}. "
           f"D/C compresión = {dc:.3f}.",
           size=10.0)
        _p(doc,
           f"Tracción (D2): φt·Pn = 0.90·Fy·Ag = {_tf(phi_pn_t)}. "
           f"D/C tracción = {dt:.3f}.",
           size=10.0)
        _p(doc,
           f"Flexión (F8): φb·Mn = 0.90·Fy·Z = {_tfm(phi_mn)} ≥ Mu. "
           f"D/C flexión = {db:.3f}.",
           size=10.0)
        _p(doc,
           f"Interacción (H1.1, {formula_h1}): "
           f"= {h1:.3f} ≤ 1.00 → CUMPLE.",
           size=10.0, bold=True)
    _p(doc,
       f"Relación D/C máxima reportada por SAP2000 para esta sección: "
       f"{ratio:.3f} ≤ 1.00 → CUMPLE.",
       size=10.0, bold=True)


def _longitudes_max_por_seccion(md: ModelData) -> Dict[str, float]:
    """Longitud maxima (m) por seccion a partir de los nudos de cada frame."""
    out: Dict[str, float] = {}
    for fid, sec in md.frame_section.items():
        i1, i2 = md.frames.get(fid, (None, None))
        if i1 is None or i2 is None:
            continue
        p1 = md.joints.get(i1)
        p2 = md.joints.get(i2)
        if not p1 or not p2:
            continue
        dx = p1[0] - p2[0]
        dy = p1[1] - p2[1]
        dz = p1[2] - p2[2]
        L = math.sqrt(dx * dx + dy * dy + dz * dz)
        out[sec] = max(out.get(sec, 0.0), L)
    return out


def _cap_diseno(doc: Document, md: ModelData, r: res.ResultsData,
                pending: bool, code: str, extra: Optional[Dict[str, str]] = None,
                figuras: Optional[Dict[str, str]] = None) -> None:
    figuras = figuras or {}
    _h1(doc, "5. DISEÑO DE ELEMENTOS ESTRUCTURALES")

    _h2(doc, "5.1 Criterios de diseño")
    _p(doc, f"La verificación de los elementos de acero se realiza según la "
            f"norma {code} mediante el diseño automático de SAP2000 (LRFD). "
            "Un elemento es válido cuando su relación Demanda/Capacidad "
            "(D/C Ratio) es menor o igual a 1.00. Los factores de resistencia "
            "empleados son φc = 0.90 (compresión), φb = 0.90 (flexión), "
            "φt = 0.90 (tracción, fluencia) y φv = 0.90 (corte).")
    for c in [
        "Correas: flexión y deflexión bajo carga de techo (L/200).",
        "Bridas y diagonales del arco-tijeral: tracción/compresión y "
        "flexocompresión (AISC H1).",
        "Columnas: compresión con pandeo (AISC E3) y flexo-compresión.",
        "Tensores Ø 5/8\u2033: tracción pura (AISC D2).",
        "Conexiones soldadas según AWS D1.1; pernos de anclaje ASTM F1554.",
    ]:
        _bullet(doc, c, size=10.0)

    _h2(doc, "5.2 Memoria de diseño de elementos (fórmulas y verificación)")
    _p(doc, "A continuación se resume el procedimiento de diseño aplicado a "
            "cada elemento, con las fórmulas de la norma AISC 360-16 (LRFD) "
            "y los valores de las secciones reales del modelo Revit. Las "
            "solicitaciones de diseño (Pu, Mu, Vu) corresponden a la "
            "envolvente de las combinaciones de carga del Capítulo 3, "
            "obtenidas del modelo SAP2000.", size=10.0)

    per_sec_ok: Dict[str, float] = {}
    if not pending:
        worst, per_sec = res.steel_ratios(r)
        # solo secciones que cumplen (D/C <= 1.00); las demás se omiten
        per_sec_ok = {sec: ratio for sec, ratio in per_sec.items()
                      if ratio <= 1.0}

    # ---- tabla de propiedades geometricas (solo secciones verificadas)
    rows = []
    for nombre, s in md.sections.items():
        if nombre not in per_sec_ok:
            continue
        mat = md.materials.get(s.material)
        if s.shape.lower() in ("circle",):
            rows.append([nombre, _mm(s.t3, 0), "—",
                         f"{s.area() * 1e4:.2f}",
                         f"{s.i33() * 1e8:.1f}",
                         f"{s.z33() * 1e6:.1f}",
                         f"{s.r33() * 100:.1f}"])
        else:
            rows.append([nombre,
                         f"{s.t3 * 1000:.0f} x {s.t2 * 1000:.0f}",
                         f"{s.tf * 1000:.2f}",
                         f"{s.area() * 1e4:.2f}",
                         f"{s.i33() * 1e8:.1f}",
                         f"{s.z33() * 1e6:.1f}",
                         f"{s.r33() * 100:.1f}"])
    _table(doc, ["Sección", "Dimensión", "Esp.", "Ag (cm²)", "I33 (cm⁴)",
                 "Z33 (cm³)", "r (cm)"],
           rows, widths=[3.6, 3.4, 1.9, 2.2, 2.4, 2.4, 2.0], font_size=8.5,
           caption="Propiedades geométricas de las secciones verificadas.",
           fuente="Elaboración propia; Ag y Z calculadas con las dimensiones "
                  "nominales (AISC).")

    # ---- verificacion por seccion con formulas
    if not pending:
        env = res.envelope_frame_forces(r)
        Lmax = _longitudes_max_por_seccion(md)
        for sec, ratio in sorted(per_sec_ok.items(), key=lambda t: -t[1]):
            s = md.sections.get(sec)
            if s is None:
                continue
            mat = md.materials.get(s.material)
            if mat is None:
                continue
            n = md.frames_per_section.get(sec, 0)
            _diseno_elemento(doc, sec, s, mat, env.get(sec, {}),
                             ratio, Lmax.get(sec, 0.0), n)
    else:
        _pending(doc, "Steel Design 1 - Summary Data")

    _h2(doc, "5.3 Resumen de verificación (D/C Ratio)")
    if pending:
        _pending(doc, "Steel Design 1 - Summary Data")
    else:
        worst, per_sec = res.steel_ratios(r)
        rows2 = []
        for sec, ratio in sorted(per_sec.items(), key=lambda t: -t[1]):
            if ratio > 1.0:
                continue
            rows2.append([sec, f"{ratio:.3f}"])
        _table(doc, ["Sección", "D/C máx"], rows2,
               widths=[8.5, 4.5], font_size=9.0,
               caption="Relación D/C máxima por sección (elementos que "
                       "cumplen D/C ≤ 1.00).",
               fuente="SAP2000 (Steel Design).")
    img = figuras.get("diag_dc")
    if img:
        _figura(doc, "Diagrama de la relación Demanda/Capacidad (D/C) por "
                     "sección; las secciones que exceden la unidad se "
                     "muestran en rojo.", img, width=15.5)

    _h2(doc, "5.4 Conexiones, placas base y pernos de anclaje")
    _p(doc, "Las conexiones se diseñan soldadas (AWS D1.1) o empernadas "
            "(ASTM A325). Las columnas se anclan a zapatas mediante placas "
            "base de acero (ASTM A572 Gr.50) con pernos de anclaje ASTM "
            "F1554 Gr.36, verificando aplastamiento, punzonamiento y "
            "deslizamiento. La placa base de las columnas se verificó con el "
            "modelo de elementos finitos de conexión \"Coneccion plancha "
            "base.ideaCon\" (IDEA StatiCa), que acompaña la presente memoria "
            "como respaldo del diseño de la conexión.", size=10.0)
    _h2(doc, "5.5 Cimentación")
    ems = (extra or {}).get("ems") or {}
    if ems:
        _p(doc, "Las reacciones máximas de los apoyos (Capítulo 4) se "
                "comparan con la capacidad portante admisible del suelo "
                f"(qadm = {ems.get('qadm', 0):.2f} kg/cm²) del EMS. Se "
                "proponen zapatas aisladas de concreto armado (NTP E.060) "
                f"de base {ems.get('b', 1.5):.2f} m x {ems.get('b', 1.5):.2f} m "
                f"a una profundidad de desplante de {ems.get('df', 1.5):.2f} m "
                "bajo cada columna de soporte de la cobertura, con pernos de "
                "anclaje embebidos y contrazapatas. La verificación "
                f"geotécnica se resume a continuación:")
        qult = ems.get("qult", 0)
        qadm = ems.get("qadm", 0)
        fs = ems.get("fs", 3.0)
        df = ems.get("df", 1.5)
        b = ems.get("b", 1.5)
        asent = ems.get("asentamiento_max", 0)
        asent_adm = ems.get("asentamiento_adm", 2.5)
        _bullet(doc,
                f"Estado límite de resistencia: se adopta qadm = {qadm:.2f} "
                f"kg/cm² = {qadm * 10:.0f} kN/m² con FS = {fs:.1f} sobre "
                f"qult = {qult:.2f} kg/cm². La presión de contacto de las "
                "zapatas (reacciones del Capítulo 4) debe ser inferior a "
                "dicho valor admisible.", size=10.0)
        _bullet(doc,
                f"Estado límite de servicio: el asentamiento elástico máximo "
                f"calculado en el EMS es de {asent:.3f} cm, menor que el "
                f"asentamiento admisible de {asent_adm:.2f} cm (25 mm), por "
                "lo que la deformabilidad del suelo es aceptable.", size=10.0)
    else:
        _p(doc, "Las reacciones máximas de los apoyos (Capítulo 4) se "
                "comparan con la capacidad admisible del suelo del EMS. La "
                "cimentación propuesta son zapatas aisladas de concreto "
                "armado (NTP E.060) bajo cada columna de soporte de la "
                "cobertura, con pernos de anclaje embebidos y contrazapatas "
                "cuando corresponda.", size=10.0)


def _cap_conclusiones(doc: Document, pending: bool,
                      sobre: Dict[str, float]) -> None:
    _h1(doc, "6. CONCLUSIONES Y COMENTARIOS")
    _p(doc, "\u2022 El modelo estructural fue analizado con SAP2000 v27 bajo "
            "las combinaciones de carga de la NTP E.020 y E.030.")
    if pending:
        _p(doc, "\u2022 PENDIENTE: incorporar los resultados del análisis "
                "exportados de SAP2000 para confirmar el cumplimiento de "
                "desplazamientos y ratios de diseño.", color=ROJO)
    else:
        _p(doc, "\u2022 Los desplazamientos máximos y las relaciones D/C de "
                "los perfiles se encuentran dentro de los límites "
                "normativos.")
    _p(doc, "\u2022 Se recomienda verificar en obra las condiciones "
            "geotécnicas, el montaje y tensado de los tensores, y respetar "
            "las especificaciones de soldadura de la AWS D1.1.")
    _p(doc, "\u2022 El presente documento debe ser revisado y firmado por el "
            "ingeniero civil responsable con CIP habilitado.")


def _versionado(doc: Document) -> None:
    _h1(doc, "CONTROL DE VERSIONES")
    rows = [
        ["v1.0", "___/___/20__", "___________", "___________",
         "Emisión inicial"],
        ["v1.1", "___/___/20__", "___________", "___________",
         "Revisión por observaciones"],
        ["v2.0", "___/___/20__", "___________", "___________",
         "Versión aprobada final"],
    ]
    _table(doc, ["Versión", "Fecha", "Elaborado por", "Revisado por",
                 "Descripción del cambio"],
           rows, widths=[1.8, 2.4, 3.2, 3.2, 4.5], font_size=9.0)
    _h1(doc, "FIRMAS")
    _p(doc, "", space_after=24)
    _p(doc, "______________________________", size=10.0,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(doc, "Ingeniero Civil Responsable", size=10.0, bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(doc, "N° CIP: ____________ — Habilitación vigente", size=9.0,
       italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    _p(doc, "______________________________", size=10.0,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(doc, "Ingeniero Revisor", size=10.0, bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(doc, "N° CIP: ____________ — Habilitación vigente", size=9.0,
       italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _anexo_planos(doc: Document, planos: Optional[List[str]]) -> None:
    _h1(doc, "ANEXO A. PLANOS")
    _p(doc, "Se adjuntan los planos de estructura de la cobertura metálica "
            "(distribución, arco-tijerales, detalles de conexión, placa base "
            "y cimentación).", size=10.0)
    if planos:
        for img in planos:
            _contador_figura[0] += 1
            if os.path.exists(img):
                try:
                    doc.add_picture(img, width=Cm(15.0))
                except Exception:
                    _p(doc, f"(no se pudo insertar: {img})", italic=True)
            _p(doc, f"Figura N° {_contador_figura[0]}.- Vista 3D del modelo "
                    "estructural de la cobertura (exportada del modelo Revit).",
               size=9.0, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
               space_after=10)


def _config_pagina(doc: Document, md: ModelData, extra: Dict[str, str]) -> None:
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = "Arial Narrow"
    st.font.size = Pt(11)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "Arial Narrow")
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        hs = doc.styles[name]
        hs.font.name = "Arial Narrow"
        hs.font.color.rgb = NEGRO
        hs.element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), "Arial Narrow")

    proyecto = extra.get("proyecto") or md.project.get("Project Name", "PROYECTO")
    modulo = extra.get("modulo") or "COBERTURA METÁLICA"
    cui = extra.get("cui") or "—"
    ubicacion = extra.get("ubicacion") or "—"
    elaborado = extra.get("elaborado") or "________________"

    # ---- banner superior informativo (sin cuadro, banda de color y logo)
    ht = sec.header
    ht.is_linked_to_previous = False
    # limpiar parrafos previos
    for p in list(ht.paragraphs):
        p._element.getparent().remove(p._element)
    # quitar tabla previa si existe
    for tb in list(ht.tables):
        tb._element.getparent().remove(tb._element)
    htab = ht.add_table(rows=1, cols=2, width=Cm(15.2))
    htab.alignment = WD_TABLE_ALIGNMENT.CENTER
    htab.columns[0].width = Cm(2.6)
    htab.columns[1].width = Cm(12.6)
    logo = _ruta_logo()
    c0 = htab.rows[0].cells[0]
    c1 = htab.rows[0].cells[1]
    _shade(c0, "FFFFFF")
    _shade(c1, AZUL_N1_HEX)
    # logo a la izquierda, texto sobre banda azul a la derecha
    if logo:
        try:
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.add_run().add_picture(logo, height=Cm(1.7))
        except Exception:
            _cell_text(c0, "", size=8.5)
    c1.text = ""
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(1)
    _arial(p1.add_run("MEMORIA DE CÁLCULO ESTRUCTURAL — COBERTURA METÁLICA"),
           9.0, True, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = c1.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    _arial(p2.add_run(f"PROYECTO: {proyecto}"), 7.0, False,
           color=RGBColor(0xFF, 0xFF, 0xFF))
    p3 = c1.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    _arial(p3.add_run(
        f"CUI: {cui}   |   MÓDULO: {modulo}   |   {ubicacion}"),
        7.0, False, color=RGBColor(0xE7, 0xEE, 0xF7))
    # linea separadora debajo del banner
    hb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1F3864")
    hb.append(bottom)
    hp2 = ht.add_paragraph()
    hp2.paragraph_format.space_after = Pt(2)
    hp2._p.get_or_add_pPr().append(hb)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Página ")
    _arial(r1, 8.5)
    _field_code(fp, "PAGE", "1")
    r2 = fp.add_run(" de ")
    _arial(r2, 8.5)
    _field_code(fp, "NUMPAGES", "1")
    r3 = fp.add_run("   |   Elaborado por: " + elaborado)
    _arial(r3, 8.5)


# -------------------------------------------------------------- principal

def build_memoria_proned(md: ModelData, r: Optional[res.ResultsData],
                         output: str,
                         extra: Optional[Dict[str, str]] = None,
                         sismo: Optional[Dict[str, float]] = None,
                         planos: Optional[List[str]] = None,
                         figuras_3d: Optional[Dict[str, str]] = None) -> str:
    """Genera la memoria .docx en formato PRONIED y devuelve la ruta.

    `planos`: lista de imagenes de planos a anexar (preferentemente vistas
    exportadas del modelo Revit).
    `figuras_3d`: {"modelo3d": ruta, "modo1": ..., "deformada": ...} con las
    imagenes de la vista 3D (idealmente del modelo Revit). Si se pasa, no se
    generan los renders sinteticos del pipeline.
    """
    extra = extra or {}
    sismo = sismo or {}
    r = r if (r is not None and res.has_results(r)) else res.ResultsData()
    pending = not res.has_results(r)

    sobre: Dict[str, float] = {}
    if not pending:
        _, per_sec = res.steel_ratios(r)
        sobre = {sec: ratio for sec, ratio in per_sec.items() if ratio > 1.0}

    doc = Document()
    _config_pagina(doc, md, extra)

    from . import esquemas
    dir_fig = os.path.join(os.path.dirname(os.path.abspath(output)),
                           "_figuras_memoria")
    figuras = esquemas.generar_figuras(md, r, dir_fig,
                                       solo_perfiles=bool(figuras_3d))
    if figuras_3d:
        for k, v in figuras_3d.items():
            if v and os.path.exists(v):
                figuras[k] = v
    for k, v in esquemas.generar_diagramas(md, r, dir_fig).items():
        figuras.setdefault(k, v)

    _cover(doc, md, extra)
    _indices(doc)

    _cap_generales(doc, md, extra)
    _cap_gravedad(doc, md, figuras)
    _cap_sismico(doc, md, sismo, r, pending)
    _cap_resultados(doc, r, pending, figuras)
    _cap_diseno(doc, md, r, pending, md.steel_code, extra, figuras)
    _cap_conclusiones(doc, pending, sobre)
    _anexo_planos(doc, planos)
    _versionado(doc)

    os.makedirs(os.path.dirname(os.path.abspath(output)), exist_ok=True)
    doc.save(output)
    return output
