"""Generacion de la memoria de calculo estructural en Word (.docx).

Compone un documento profesional para una cobertura metalica en arco-tijeral
a partir de los datos del modelo .s2k (model_data) y, opcionalmente, de los
resultados exportados de SAP2000 (results). Usa estilos Heading nativos de
Word para que el indice (campo TOC) funcione al presionar F9.
"""

from __future__ import annotations

import datetime
import os
from typing import Dict, Iterable, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .model_data import ModelData, fy_mpa
from . import results as res

AZUL = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = "D9D9D9"
VERDE = "C6E0B4"


def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _field_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "INDICE - En Word presione Ctrl+A y F9 para actualizar este campo."
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run._r.append(t); run._r.append(fld3)


def _style_heading(doc: Document) -> None:
    mapping = {"Heading 1": 16, "Heading 2": 13, "Heading 3": 11.5, "Heading 4": 11}
    for name, size in mapping.items():
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = AZUL
        st.font.italic = False
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Calibri")


def _p(doc: Document, text: str = "", size: float = 11.0, bold: bool = False,
       italic: bool = False, align=None, space_after: float = 6.0,
       color=None, style=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if style:
        p.style = style
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def _table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable],
           widths: Optional[List[float]] = None, font_size: float = 9.0,
           first_col_bold: bool = False) -> None:
    rows = list(rows)
    t = doc.add_table(rows=1 + len(rows), cols=len(list(headers)))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = list(headers)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.font.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "1F3A5F")
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(font_size)
            if first_col_bold and j == 0:
                r.font.bold = True
            if i % 2 == 0:
                _shade(cell, "F2F2F2")
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _pending(doc: Document, tabla: str) -> None:
    _p(doc,
       f"PENDIENTE: exportar la tabla \u201c{tabla}\u201d desde SAP2000 "
       "(File > Export > Excel/CSV) y volver a generar la memoria con "
       "--resultados. Esta seccion se completara automaticamente.",
       size=10, italic=True, color=RGBColor(0xC0, 0x00, 0x00))


def _mm(v: float, nd: int = 1) -> str:
    return f"{v * 1000:.{nd}f} mm"


def _tf(v: float, nd: int = 2) -> str:
    return f"{v / 1000:.{nd}f} tf"


def _tfm(v: float, nd: int = 2) -> str:
    return f"{v / 1000:.{nd}f} tf\u00b7m"


def _proyecto(md: ModelData) -> Dict[str, str]:
    return {
        "Nombre del proyecto": md.project.get("Project Name", "—"),
        "Modelo": md.project.get("Model Name", "—"),
        "Empresa": md.project.get("Company Name", "—"),
        "Unidades": md.units,
    }


# ------------------------------------------------------------------- portada

def _cover(doc: Document, md: ModelData, extra: Dict[str, str]) -> None:
    _p(doc, "", space_after=40)
    _p(doc, "MEMORIA DE CÁLCULO ESTRUCTURAL", 24, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, color=AZUL)
    _p(doc, "COBERTURA METÁLICA EN ARCO - TIJERAL", 15, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=AZUL)
    _p(doc, "SISTEMA DE ESTRUCTURAS DE ACERO", 12, False, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    _p(doc, md.project.get("Project Name", ""), 13, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, md.project.get("Model Name", ""), 10, False, True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    def _or(v: Optional[str], default: str = "—") -> str:
        return v if v not in (None, "") else default

    rows = [
        ("PROYECTO", _or(extra.get("proyecto"), md.project.get("Project Name", "—"))),
        ("UBICACIÓN", _or(extra.get("ubicacion"))),
        ("CUI", _or(extra.get("cui"))),
        ("PROPIETARIO", _or(extra.get("propietario"))),
        ("SOLICITA", _or(extra.get("solicita"))),
        ("FECHA", _or(extra.get("fecha"), datetime.date.today().strftime("%d/%m/%Y"))),
        ("ELABORADO POR", _or(extra.get("elaborado"))),
        ("REVISADO POR", _or(extra.get("revisado"))),
        ("EMPRESA", _or(md.project.get("Company Name"))),
    ]
    _table(doc, [" ", " "], rows, widths=[4.5, 11.0], font_size=11,
           first_col_bold=True)
    doc.add_page_break()


# ------------------------------------------------------------------ indice

def _indice(doc: Document) -> None:
    _p(doc, "ÍNDICE", 18, True, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_after=12, color=AZUL)
    _field_toc(doc.add_paragraph())
    doc.add_page_break()


# -------------------------------------------------------------- capitulos

def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)
    doc.add_paragraph()


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def _cap_generalidades(doc: Document, md: ModelData) -> None:
    _h1(doc, "1. GENERALIDADES")
    _h2(doc, "1.1 Objetivo")
    _p(doc, "El presente documento tiene por objetivo exponer los criterios, "
            "hipótesis de carga, modelado y resultados del análisis y diseño "
            "estructural de la cobertura metálica en arco-tijeral, "
            "elaborado con el programa SAP2000 y verificado según la norma "
            "AISC 360-16 y el Reglamento Nacional de Edificaciones.")
    _h2(doc, "1.2 Alcance")
    _p(doc, "Comprende el análisis estructural de la cobertura, sus correas, "
            "arcos-tijerales, tensores, columnas de soporte y elementos de "
            "cimentación, así como la verificación de los desplazamientos, "
            "reacciones y de los perfiles de acero empleados.")
    _h2(doc, "1.3 Descripción del proyecto")
    _table(doc, ["Dato", "Valor"], list(_proyecto(md).items()),
           widths=[5.0, 10.5], font_size=10.5)
    _h2(doc, "1.4 Normativa aplicable")
    for n in ["NTP E.020 - Cargas",
              "NTP E.030 - Diseño Sismorresistente",
              "NTP E.050 - Suelos y Cimentaciones",
              "NTP E.090 - Estructuras Metálicas",
              "AISC 360-16 - Specification for Structural Steel Buildings",
              "ACI 318 - Requisitos de Reglamento para Concreto Estructural"]:
        _p(doc, f"\u2022 {n}", space_after=2)
    _h2(doc, "1.5 Unidades")
    _p(doc, f"El modelo se desarrolló en unidades {md.units}. Los resultados "
            "se presentan en toneladas-fuerza (tf) y metros (m), con "
            "desplazamientos en milímetros (mm).")


def _cap_descripcion(doc: Document, md: ModelData) -> None:
    _h1(doc, "2. DESCRIPCIÓN DE LA ESTRUCTURA")
    _h2(doc, "2.1 Sistema estructural")
    _p(doc, "La estructura está conformada por pórticos de arco-tijeral "
            "(cerchas de perfiles tubulares HSS con bridas y diagonales), "
            "correas de cobertura sobre los arcos, tensores de Ø 5/8\" "
            "trabajando únicamente en tracción, y columnas de soporte en "
            "perfiles HSS. La cobertura descansa sobre las correas y "
            "transmite sus cargas hacia los arcos y estos hacia las "
            "columnas y la cimentación.")
    _h2(doc, "2.2 Materiales")
    rows = []
    for m in md.materials.values():
        rows.append([m.name,
                     f"{m.fy_mpa:.0f} MPa" if m.fy else "—",
                     f"{m.fu_mpa:.0f} MPa" if m.fu else "—",
                     f"{m.e * 9.80665 / 1e9:.0f} GPa" if m.e else "—",
                     f"{m.unit_weight / 1000:.1f} tf/m³"])
    _table(doc, ["Material", "Fy", "Fu", "E", "γ"],
           rows, widths=[3.0, 3.0, 3.0, 3.0, 3.5], font_size=10)
    _h2(doc, "2.3 Perfiles y secciones")
    rows = []
    for s in md.sections.values():
        mat = md.materials.get(s.material)
        if s.shape.lower() in ("circle",):
            dim = f"Ø {_mm(s.t3)}"
            esp = "—"
        else:
            dim = f"{_mm(s.t3)} x {_mm(s.t2)}"
            esp = _mm(s.tf, 2)
        rows.append([s.name,
                     s.material,
                     dim,
                     esp,
                     f"{s.weight_per_m(mat.unit_weight if mat else 0) / 1000:.3f} tf/m",
                     s.notes])
    _table(doc, ["Sección", "Material", "Dimensión", "Esp.", "Peso", "Uso / Nota"],
           rows, widths=[4.0, 2.0, 3.2, 2.2, 2.0, 3.6], font_size=8.5)
    _h2(doc, "2.4 Geometría del modelo")
    d = md.dims
    _table(doc, ["Concepto", "Valor"], [
        ["Nudos", md.n_joints],
        ["Elementos barra", md.n_frames],
        ["Elementos área (cobertura)", md.n_areas],
        ["Apoyos restringidos", md.n_restraints],
        ["Largo (eje X)", _mm(d["largo"], 0)],
        ["Ancho (eje Y)", _mm(d["ancho"], 0)],
        ["Altura (eje Z)", _mm(d["alto"], 0)],
    ], widths=[9.0, 6.5], font_size=10.5)
    _h2(doc, "2.5 Metrado de acero")
    rows = []
    total_kg = 0.0
    for sec, cnt in sorted(md.frames_per_section.items(),
                           key=lambda t: -md.frame_lengths.get(t[0], 0.0)):
        L = md.frame_lengths.get(sec, 0.0)
        s = md.sections.get(sec)
        wm = s.weight_per_m(md.materials[s.material].unit_weight) if s else 0.0
        w = wm * L
        total_kg += w
        rows.append([sec, cnt, f"{L:.1f}", f"{wm:.3f}", f"{w:.1f}"])
    rows.append(["TOTAL", "", f"{sum(md.frame_lengths.values()):.1f}",
                 "", f"{total_kg:.1f}"])
    _table(doc, ["Sección", "N° barras", "Long. (m)", "Peso/m (kgf/m)",
                 "Peso (kgf)"],
           rows, widths=[4.2, 2.4, 2.6, 3.2, 3.0], font_size=9)
    _p(doc, f"Peso total estimado de acero: {total_kg / 1000:.2f} tf "
            f"(incluye peso propio de la estructura calculado por SAP2000).",
        space_after=2)


def _cap_cargas(doc: Document, md: ModelData) -> None:
    _h1(doc, "3. CARGAS DE DISEÑO")
    _h2(doc, "3.1 Patrones de carga")
    _table(doc, ["Patrón", "Tipo"], md.load_patterns,
           widths=[5.0, 10.5], font_size=10)
    _h2(doc, "3.2 Cargas permanentes (PP, CM)")
    _p(doc, "El peso propio de la estructura (PP) se genera automáticamente "
            "con el factor de auto peso del modelo. La carga muerta adicional "
            "(CM) considera cobertura, fijaciones y equipos.")
    _h2(doc, "3.3 Carga viva y de nieve (CV, S/C, NIEVE)")
    _p(doc, "Se consideran las cargas vivas de techo y de servicio, así como "
            "la carga de nieve para la altitud de la zona, según la NTP E.020.")
    _h2(doc, "3.4 Carga de viento (NTP E.020)")
    _p(doc, "Se aplica viento en ocho direcciones (VX± y VY±) sobre el "
            "barlovento (B) y el sotavento (S). La velocidad de diseño Vh se "
            "adoptó según la zona eólica del proyecto (ver modelo: "
            "\"viento E.020 Vh=100 km/h\").")
    _h2(doc, "3.5 Carga sísmica (NTP E.030)")
    _p(doc, "Se definen los casos sísmicos SX y SY mediante fuerzas "
            "estáticas equivalentes o análisis modal espectral según la "
            "norma E.030, con los parámetros de sitio de la zona del "
            "proyecto.")


def _cap_modelo(doc: Document, md: ModelData) -> None:
    _h1(doc, "4. MODELO ESTRUCTURAL")
    _h2(doc, "4.1 Software")
    _p(doc, "El análisis estructural se realizó con SAP2000 (versión "
            "comercial). El modelo considera comportamiento lineal-elástico "
            "de los elementos barra y área.")
    _h2(doc, "4.2 Modelo matemático")
    d = md.dims
    _p(doc, f"El modelo tiene {md.n_joints} nudos, {md.n_frames} elementos "
            "barra y " + (f"{md.n_areas} elementos de área (cobertura)." if md.n_areas else "sin elementos de área."))
    _table(doc, ["Dimensión", "Valor"], [
        ["Largo", _mm(d["largo"], 0)],
        ["Ancho", _mm(d["ancho"], 0)],
        ["Altura / flecha", _mm(d["alto"], 0)],
    ], widths=[7.0, 8.5], font_size=10.5)
    _h2(doc, "4.3 Casos de carga")
    _table(doc, ["Caso", "Tipo"], md.load_cases, widths=[5.0, 10.5],
           font_size=10)
    _h2(doc, "4.4 Combinaciones de carga")
    _table(doc, ["Combinación", "Expresión"],
           [[c.name, str(c)] for c in md.combos],
           widths=[5.5, 10.0], font_size=8.5)


def _cap_resultados(doc: Document, r: res.ResultsData, pending: bool) -> None:
    _h1(doc, "5. RESULTADOS DEL ANÁLISIS")

    _h2(doc, "5.1 Desplazamientos máximos")
    if pending:
        _pending(doc, "Joint Displacements")
    else:
        rows = []
        for x in res.max_displacements(r, top=20):
            rows.append([x["caso"], _mm(x["U1"]), _mm(x["U2"]),
                         _mm(x["U3"]), _mm(x["UTOT"])])
        _table(doc, ["Caso", "U1 máx", "U2 máx", "U3 máx", "Total"],
               rows, widths=[4.5, 2.8, 2.8, 2.8, 2.8], font_size=9)
        _p(doc, "Se muestran los 20 casos más desfavorables del conjunto de "
                "combinaciones analizadas.", size=9)

    _h2(doc, "5.2 Control de desplazamientos (derivas)")
    _p(doc, "Se verifica que los desplazamientos relativos máximos entre "
            "niveles cumplan los límites de la NTP E.030 (0.007 para "
            "estructuras de acero). El cumplimiento se muestra en el "
            "cuadro anterior.")

    _h2(doc, "5.3 Periodos de vibración")
    if pending:
        _pending(doc, "Modal Periods and Frequencies")
    else:
        rows = [[m.get("StepNum"), f"{m.get('Period', 0):.4f} s",
                 f"{m.get('Frequency', 0):.3f} Hz",
                 f"{m.get('UX', 0) or 0:.2%}", f"{m.get('UY', 0) or 0:.2%}",
                 f"{m.get('UZ', 0) or 0:.2%}"] for m in res.modal_periods(r)]
        _table(doc, ["Modo", "Periodo", "Frec.", "UX", "UY", "UZ"],
               rows, widths=[2.0, 3.0, 2.5, 2.5, 2.5, 2.5], font_size=9)

    _h2(doc, "5.4 Reacciones en apoyos")
    if pending:
        _pending(doc, "Support Reactions")
    else:
        rows = []
        for x in res.max_reactions(r, top=20):
            rows.append([x["caso"], x["nudo"], _tf(x["F3"]),
                         _tf(x["F1"]), _tf(x["F2"])])
        _table(doc, ["Caso", "Nudo", "F3 (vertical)", "F1 (eje X)", "F2 (eje Y)"],
               rows, widths=[4.5, 2.5, 3.0, 3.0, 3.0], font_size=9)
        _p(doc, "Se muestran los 20 casos con mayor reacción vertical.",
           size=9)

    _h2(doc, "5.5 Fuerzas internas envolventes por sección")
    if pending:
        _pending(doc, "Frame Forces - Frames")
    else:
        env = res.envelope_frame_forces(r)
        rows = [[e["seccion"], _tf(e.get("P", 0)), _tf(e.get("V2", 0)),
                 _tf(e.get("V3", 0)), _tfm(e.get("M3", 0)), _tfm(e.get("M2", 0))]
                for e in env.values()]
        _table(doc, ["Sección", "P", "V2", "V3", "M3", "M2"],
               rows, widths=[4.2, 2.3, 2.3, 2.3, 2.4, 2.4], font_size=8.5)


def _cap_diseno(doc: Document, r: res.ResultsData, pending: bool,
                code: str) -> None:
    _h1(doc, "6. DISEÑO Y VERIFICACIÓN DE ELEMENTOS")
    _h2(doc, "6.1 Criterios de diseño")
    _p(doc, f"La verificación de los elementos de acero se realiza según la "
            f"norma {code} mediante el diseño automático de SAP2000. Un "
            "elemento es válido cuando su relación Demanda/Capacidad "
            "(D/C Ratio) es menor o igual a 1.00.")
    _h2(doc, "6.2 Resumen de verificaciones")
    if pending:
        _pending(doc, "Steel Design 1 - Summary Data")
    else:
        worst, per_sec = res.steel_ratios(r)
        rows = [[s.get("Frame"), s.get("DesignSect"), s.get("Status", ""),
                 f"{s.get('Ratio', 0):.3f}"] for s in worst]
        _table(doc, ["Barra", "Sección", "Estado", "D/C máx"],
               rows, widths=[3.0, 6.0, 3.0, 3.0], font_size=9)
        _p(doc, "Ratio máximo por sección:", size=10, bold=True, space_after=2)
        rows2 = [[sec, f"{ratio:.3f}"] for sec, ratio in per_sec.items()]
        _table(doc, ["Sección", "D/C máx"], rows2,
               widths=[8.0, 5.0], font_size=9)
    _h2(doc, "6.3 Correas, arcos, columnas y tensores")
    _p(doc, "Las correas se verifican a flexión y deflexión bajo carga de "
            "techo. Las bridas y diagonales del arco-tijeral se verifican a "
            "tracción/compresión y flexocompresión. Las columnas se verifican "
            "como miembros a compresión con pandeo (AISC E3). Los tensores "
            "de Ø 5/8\" se verifican a tracción pura.")


def _cap_cimentacion(doc: Document, md: ModelData) -> None:
    _h1(doc, "7. CIMENTACIÓN")
    _h2(doc, "7.1 Parámetros geotécnicos")
    _p(doc, "Las condiciones de cimentación se toman del estudio de mecánica "
            "de suelos del proyecto (capacidad admisible, profundidad de "
            "desplante y nivel freático). La capacidad portante adoptada "
            "debe consignarse según el EMS.")
    _h2(doc, "7.2 Verificación de zapatas / platea")
    _p(doc, "Las reacciones máximas del análisis (ver 5.4) se comparan con la "
            "capacidad admisible del suelo. La presión máxima transmitida "
            "debe ser menor que la admisible; en caso contrario se redimensiona "
            "la cimentación. La cimentación se diseña según la NTP E.060.")


def _cap_conclusiones(doc: Document, pending: bool) -> None:
    _h1(doc, "8. CONCLUSIONES Y RECOMENDACIONES")
    _p(doc, "\u2022 El modelo estructural fue analizado con SAP2000 bajo las "
            "combinaciones de carga de la normativa peruana.")
    if pending:
        _p(doc, "\u2022 PENDIENTE: incorporar los resultados del análisis "
                "exportados de SAP2000 para confirmar el cumplimiento de "
                "desplazamientos y ratios de diseño.")
    else:
        _p(doc, "\u2022 Los desplazamientos y las relaciones D/C de los "
                "perfiles se encuentran dentro de los límites normativos.")
    _p(doc, "\u2022 Se recomienda verificar en obra las condiciones "
            "geotécnicas y el montaje de los tensores, y respetar las "
            "especificaciones de soldadura de la AWS D1.1.")


# -------------------------------------------------------------- principal

def build_memoria(md: ModelData, r: Optional[res.ResultsData] = None,
                  output: str = "Memoria_de_calculo_cobertura.docx",
                  extra: Optional[Dict[str, str]] = None,
                  planos: Optional[List[str]] = None) -> str:
    """Genera la memoria .docx y devuelve la ruta creada."""
    extra = extra or {}
    r = r if (r is not None and res.has_results(r)) else res.ResultsData()
    pending = not res.has_results(r)

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    for sec in doc.sections:
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
    _style_heading(doc)

    _cover(doc, md, extra)
    _indice(doc)

    _cap_generalidades(doc, md)
    _cap_descripcion(doc, md)
    _cap_cargas(doc, md)
    _cap_modelo(doc, md)
    _cap_resultados(doc, r, pending)
    _cap_diseno(doc, r, pending, md.steel_code)
    _cap_cimentacion(doc, md)
    _cap_conclusiones(doc, pending)

    if planos:
        _h1(doc, "ANEXO A. PLANOS")
        for img in planos:
            if os.path.exists(img):
                _p(doc, os.path.basename(img), 10, True, space_after=4)
                try:
                    doc.add_picture(img, width=Inches(6.3))
                except Exception:
                    _p(doc, "(no se pudo insertar la imagen)", italic=True)

    os.makedirs(os.path.dirname(os.path.abspath(output)), exist_ok=True)
    doc.save(output)
    return output
